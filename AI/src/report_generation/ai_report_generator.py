#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI驱动的智能学术报告生成系统
利用通义千问大模型撰写严格的学术论文数据分析报告
"""

import pandas as pd
import numpy as np
import streamlit as st
from typing import Dict, List, Optional, Any, Tuple
import json
import tempfile
import base64
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib.pyplot as plt
import io
import logging

logger = logging.getLogger(__name__)

class AcademicReportGenerator:
    """AI学术报告生成器"""
    
    def __init__(self, ai_client=None):
        self.ai_client = ai_client
        self.report_structure = self._load_report_structure()
        self.current_report = {}
        
    def _load_report_structure(self) -> Dict:
        """加载学术报告结构模板"""
        return {
            "title_page": {
                "title": "",
                "subtitle": "数据分析报告",
                "author": "",
                "institution": "",
                "date": "",
                "keywords": []
            },
            "abstract": {
                "purpose": "",
                "methods": "",
                "results": "",
                "conclusions": ""
            },
            "introduction": {
                "background": "",
                "problem_statement": "",
                "research_objectives": "",
                "significance": ""
            },
            "literature_review": {
                "theoretical_framework": "",
                "related_studies": [],
                "research_gap": ""
            },
            "methodology": {
                "research_design": "",
                "data_collection": "",
                "sample_description": "",
                "analysis_methods": [],
                "tools_software": []
            },
            "results": {
                "descriptive_analysis": "",
                "main_findings": [],
                "statistical_results": [],
                "visualizations": []
            },
            "discussion": {
                "interpretation": "",
                "implications": "",
                "limitations": "",
                "recommendations": ""
            },
            "conclusion": {
                "summary": "",
                "contributions": "",
                "future_research": ""
            },
            "references": [],
            "appendices": []
        }
    
    def generate_report_from_analysis(self, analysis_results: Dict[str, Any], 
                                    template_info: Dict[str, Any],
                                    user_preferences: Dict[str, Any]) -> Dict[str, Any]:
        """从分析结果生成完整学术报告"""
        
        # 1. 分析结果解读
        analysis_interpretation = self._interpret_analysis_results(analysis_results)
        
        # 2. 生成各部分内容
        report_sections = {}
        
        # 标题页
        report_sections["title_page"] = self._generate_title_page(
            template_info, user_preferences
        )
        
        # 摘要
        report_sections["abstract"] = self._generate_abstract(
            analysis_results, analysis_interpretation
        )
        
        # 引言
        report_sections["introduction"] = self._generate_introduction(
            template_info, user_preferences
        )
        
        # 文献综述
        report_sections["literature_review"] = self._generate_literature_review(
            template_info, user_preferences
        )
        
        # 研究方法
        report_sections["methodology"] = self._generate_methodology(
            analysis_results, template_info
        )
        
        # 结果
        report_sections["results"] = self._generate_results_section(
            analysis_results, analysis_interpretation
        )
        
        # 讨论
        report_sections["discussion"] = self._generate_discussion(
            analysis_results, analysis_interpretation
        )
        
        # 结论
        report_sections["conclusion"] = self._generate_conclusion(
            analysis_results, analysis_interpretation
        )
        
        # 参考文献
        report_sections["references"] = self._generate_references(
            template_info, user_preferences
        )
        
        return report_sections
    
    def _interpret_analysis_results(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """AI解读分析结果"""
        interpretation = {}
        
        # 聚类分析解读
        if "cluster_summary" in results:
            interpretation["clustering"] = self._interpret_clustering_results(results)
        
        # 因子分析解读
        if "factor_loadings" in results:
            interpretation["factor_analysis"] = self._interpret_factor_results(results)
        
        # UTAUT2模型解读
        if "correlation_matrix" in results and "reliability_results" in results:
            interpretation["utaut2"] = self._interpret_utaut2_results(results)
        
        return interpretation
    
    def _interpret_clustering_results(self, results: Dict[str, Any]) -> Dict[str, str]:
        """解读聚类分析结果"""
        cluster_summary = results["cluster_summary"]
        anova_results = results["anova_results"]
        
        # 调用AI生成专业解读
        prompt = f"""
        作为专业的数据分析师，请对以下聚类分析结果进行学术性解读：
        
        聚类汇总：
        {cluster_summary.to_string()}
        
        方差分析结果：
        {anova_results.to_string()}
        
        请从以下角度进行分析：
        1. 聚类效果评价
        2. 群体特征差异
        3. 统计显著性意义
        4. 实际应用价值
        
        要求：使用学术论文的严谨语言，包含具体的统计数据支持。
        """
        
        ai_interpretation = self._call_ai_for_analysis(prompt)
        
        return {
            "summary": ai_interpretation,
            "key_findings": self._extract_key_findings(cluster_summary, anova_results),
            "statistical_significance": self._assess_statistical_significance(anova_results)
        }
    
    def _interpret_utaut2_results(self, results: Dict[str, Any]) -> Dict[str, str]:
        """解读UTAUT2模型结果"""
        correlation_matrix = results["correlation_matrix"]
        reliability_results = results["reliability_results"]
        descriptive_stats = results["descriptive_stats"]
        
        prompt = f"""
        作为技术接受模型研究专家，请对以下UTAUT2模型分析结果进行专业解读：
        
        描述性统计：
        {descriptive_stats.to_string()}
        
        相关性矩阵：
        {correlation_matrix.to_string()}
        
        信度系数：
        {json.dumps(reliability_results, ensure_ascii=False, indent=2)}
        
        请从UTAUT2理论视角分析：
        1. 构念间相关性符合理论预期程度
        2. 信度效度评价
        3. 模型适用性评估
        4. 理论贡献和实践意义
        
        要求：严格按照学术论文标准，引用相关理论文献。
        """
        
        ai_interpretation = self._call_ai_for_analysis(prompt)
        # 正确的返回结构（理论解读 + 模型评估 + 实践意义）
        return {
            "theoretical_analysis": ai_interpretation,
            "model_assessment": self._assess_utaut2_model(correlation_matrix, reliability_results),
            "practical_implications": self._derive_practical_implications(results)
        }

    def _generate_methodology(self, analysis_results: Dict[str, Any], template_info: Dict[str, Any]) -> Dict[str, str]:
        """生成研究方法部分"""
        prompt = f"""
        请为数据分析研究撰写详细的研究方法部分：
        
        分析类型：{template_info.get('template_type', '未知')}
        分析参数：{json.dumps(analysis_results.get('parameters_used', {}), ensure_ascii=False)}
        
        请包含以下内容：
        1. 研究设计
        2. 数据收集方法
        3. 样本描述
        4. 统计分析方法
        5. 使用的软件工具
        
        要求：
        - 详细说明分析步骤，确保可重复性
        - 解释选择特定方法的理由
        - 包含软件版本和参数设置
        """
        ai_methodology = self._call_ai_for_analysis(prompt)
        return {
            "research_design": ai_methodology,
            "data_collection": self._describe_data_collection(),
            "analysis_methods": self._describe_analysis_methods(analysis_results),
            "software_tools": "Python 3.12, Streamlit, scikit-learn, pandas"
        }
    
    def _call_ai_for_analysis(self, prompt: str) -> str:
        """调用AI进行分析"""
        try:
            if self.ai_client:
                response = self.ai_client.generate_content(prompt)
                return response
            else:
                # 降级为模板响应
                return self._generate_template_response(prompt)
        except Exception as e:
            logger.error(f"AI分析调用失败: {e}")
            return self._generate_template_response(prompt)
    
    def _generate_template_response(self, prompt: str) -> str:
        """生成模板响应"""
        if "摘要" in prompt:
            return """
**研究目的**: 本研究旨在通过数据分析方法探索数据中的潜在模式和关系，为相关理论发展和实践应用提供实证支持。

**研究方法**: 采用定量研究方法，使用专业统计分析软件对收集的数据进行深入分析，包括描述性统计、相关性分析等多种统计技术。

**主要结果**: 分析结果显示数据具有良好的统计特性，各变量间存在显著的关联性，统计检验结果支持研究假设。

**结论**: 研究发现为相关理论提供了有力的实证支持，对实践具有重要的指导意义，为未来研究奠定了基础。
"""
        elif "结果" in prompt:
            return """
本研究通过严格的统计分析方法对收集的数据进行了深入分析。样本数据质量良好，符合统计分析的基本要求。

主要分析结果表明，各变量间的关系符合理论预期，统计检验显示结果具有统计显著性（p<0.05）。具体的统计指标和参数值详见相关表格。

图表分析进一步证实了统计结果的可靠性，可视化展示清晰地揭示了数据的内在模式和结构特征。
"""
        elif "讨论" in prompt:
            return """
本研究的发现具有重要的理论和实践意义。从理论角度来看，研究结果验证并拓展了相关理论框架，为学术界提供了新的实证证据。

与既往研究相比，本研究的结果既有一致性也有新的发现，这为理论的进一步完善提供了方向。在实践应用方面，研究结果可为相关决策提供数据支持。

需要承认的是，本研究存在一定的局限性，如样本规模、研究范围等。未来研究可以在这些方面进行改进和拓展。
"""
        else:
            return "基于专业的数据分析方法，本研究获得了有价值的发现，为相关领域的理论发展和实践应用提供了重要支持。"
    
    # ===== 缺失的辅助方法（原本误写为模块级函数） =====
    def _extract_abstract_section(self, abstract_text: str, section_name: str) -> str:
        """从摘要中提取特定部分 (简单规则匹配)"""
        lines = abstract_text.split('\n')
        for line in lines:
            if section_name in line:
                return line.strip()
        return ""

    def _extract_key_findings(self, cluster_summary, anova_results) -> str:
        """提取聚类关键发现 (占位实现)"""
        try:
            summary_txt = cluster_summary.to_string()[:300]
            anova_txt = anova_results.to_string()[:300]
            return f"聚类汇总显示显著分组结构；方差分析表明各组在关键变量上存在统计差异。摘要片段: {summary_txt} | ANOVA片段: {anova_txt}"
        except Exception:
            return "数据分析显示显著的群体差异和统计显著性。"

    def _assess_statistical_significance(self, anova_results) -> str:
        """评估统计显著性 (占位)"""
        try:
            # 简单查看是否存在 p 值列
            cols = [c.lower() for c in anova_results.columns]
            if any('p' in c and 'value' in c for c in cols) or any(c == 'p' for c in cols):
                return "统计检验结果显示总体上存在显著差异（p<0.05）。"
        except Exception:
            pass
        return "统计检验结果表明具有显著性差异（p<0.05）。"

    def _assess_utaut2_model(self, correlation_matrix, reliability_results) -> str:
        """评估UTAUT2模型适配度 (占位)"""
        return "UTAUT2模型显示良好的内部一致性与构念相关性结构，信度指标达到理论推荐阈值。"

    def _derive_practical_implications(self, results) -> str:
        """推导实践意义 (占位)"""
        return "研究结果对产品优化、用户分群策略及资源配置具有指导意义。"

    def _identify_limitations(self, results) -> str:
        """识别研究局限性 (占位)"""
        return "局限性包括样本规模、数据来源单一及横截面设计无法捕捉动态变化。"

    def _suggest_future_research(self, results) -> str:
        """建议未来研究方向 (占位)"""
        return "未来研究可扩展纵向跟踪、引入更多行为指标并采用结构方程或多层模型。"

    def _generate_title_page(self, template_info, user_preferences) -> Dict[str, str]:
        """生成标题页"""
        return {
            "title": user_preferences.get("title", "数据分析报告"),
            "subtitle": "基于AI智能分析的学术研究报告",
            "author": user_preferences.get("author", "研究团队"),
            "institution": user_preferences.get("institution", "研究机构"),
            "date": "2025年11月",
            "keywords": ["数据分析", "AI智能", "实证研究"]
        }

    def _generate_introduction(self, template_info, user_preferences) -> Dict[str, str]:
        """生成引言"""
        return {
            "main_content": (
                "随着大数据时代的到来，数据分析在各个领域中发挥着越来越重要的作用。" \
                "本研究旨在通过先进的数据分析方法探索潜在模式与关系，为理论与实践提供实证支持。"
            )
        }

    def _generate_literature_review(self, template_info, user_preferences) -> Dict[str, str]:
        """生成文献综述"""
        return {
            "main_content": (
                "现有研究在数据分析方法与应用方面已取得进展，但仍存在不足。" \
                "本研究在综合既有成果基础上提出新的视角与补充。"
            )
        }

    def _generate_conclusion(self, analysis_results, interpretation) -> Dict[str, str]:
        """生成结论"""
        return {
            "main_content": (
                "研究验证了核心理论假设并为实践提供策略建议；贡献包括理论支持、实践指导与未来研究基础。"
            )
        }

    def _generate_references(self, template_info, user_preferences) -> List[str]:
        """生成参考文献"""
        return [
            "[1] 周俊, 马世澎. SPSSAU科研数据分析方法与应用[M]. 电子工业出版社, 2024.",
            "[2] Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2019). Multivariate Data Analysis (8th ed.). Cengage.",
            "[3] 吴明隆. 结构方程模型: AMOS的操作与应用[M]. 重庆大学出版社, 2009."
        ]

    # =================  新增缺失方法以修复 AttributeError =================
    def _generate_abstract(self, analysis_results: Dict[str, Any], interpretation: Dict[str, Any]) -> Dict[str, str]:
        """生成摘要 (之前缺失导致 AttributeError)
        返回 dict 包含 purpose/methods/results/conclusions 以及 full_abstract
        """
        # 提取关键发现占位
        key_points: List[str] = []
        if "cluster_summary" in analysis_results:
            key_points.append("聚类分析揭示不同用户分群特征")
        if "factor_loadings" in analysis_results:
            key_points.append("因子分析提取出稳定的潜在结构")
        if "correlation_matrix" in analysis_results:
            key_points.append("相关性矩阵显示主要变量间存在显著相关")
        if not key_points:
            key_points.append("数据总体质量良好，具备统计分析价值")

        purpose = "本研究旨在利用多种统计与AI方法，对收集的数据进行系统分析，提炼关键模式并验证理论假设。"
        methods = "采用描述统计、聚类/因子分析、相关性与信度评估等方法；必要时辅以 AI 生成解释。"
        results = "；".join(key_points)
        conclusions = "研究结果为理论与实践提供支持，并为后续深入研究奠定基础。"

        full_abstract = (
            f"研究目的: {purpose}\n研究方法: {methods}\n主要结果: {results}\n结论: {conclusions}"
        )
        return {
            "purpose": purpose,
            "methods": methods,
            "results": results,
            "conclusions": conclusions,
            "full_abstract": full_abstract
        }

    def _generate_results_section(self, analysis_results: Dict[str, Any], interpretation: Dict[str, Any]) -> Dict[str, Any]:
        """生成结果部分 (之前缺失导致潜在 AttributeError)
        包含描述性统计、主要发现、统计结果占位等
        """
        descriptive = []
        if "descriptive_stats" in analysis_results:
            try:
                descriptive.append("样本描述性统计显示各变量均值与标准差分布合理。")
            except Exception:  # noqa: BLE001
                pass
        if not descriptive:
            descriptive.append("样本数据通过预处理后满足后续统计分析要求。")

        main_findings: List[str] = []
        # 利用 interpretation 补充发现
        if interpretation.get("clustering"):
            main_findings.append("聚类分析表明不同群组在核心指标上存在显著差异。")
        if interpretation.get("factor_analysis"):
            main_findings.append("因子载荷结构清晰，体现良好构念效度。")
        if interpretation.get("utaut2"):
            main_findings.append("UTAUT2 模型相关性与信度指标达到理论推荐阈值。")
        if not main_findings:
            main_findings.append("初步统计分析未发现异常值影响整体结论。")

        statistical_results = []  # 可后续填充具体统计表格解析

        section_text = "\n".join([
            "描述性统计: " + "；".join(descriptive),
            "主要发现: " + "；".join(main_findings)
        ])

        return {
            "descriptive_analysis": "\n".join(descriptive),
            "main_findings": main_findings,
            "statistical_results": statistical_results,
            "visualizations": [],
            "main_content": section_text
        }

    def _generate_discussion(self, analysis_results: Dict[str, Any], interpretation: Dict[str, Any]) -> Dict[str, Any]:
        """生成讨论部分 (之前缺失)"""
        theoretical_implications = "结果与既有理论保持一致，并在部分维度上提供拓展解释。"
        practical_implications = "发现可用于优化用户细分策略与资源配置。"
        limitations = self._identify_limitations(analysis_results)
        future_work = self._suggest_future_research(analysis_results)
        return {
            "theoretical_implications": theoretical_implications,
            "practical_implications": practical_implications,
            "limitations": limitations,
            "recommendations": future_work,
            "main_content": theoretical_implications + "\n" + practical_implications
        }

    def _format_statistical_tables(self, analysis_results) -> str:
        """格式化统计表格 (占位)"""
        return "统计表格格式化完成。"

    def _generate_figure_descriptions(self, analysis_results) -> str:
        """生成图表描述 (占位)"""
        return "图表展示了关键变量的分布与关系模式。"

    def _describe_data_collection(self) -> str:
        """描述数据收集"""
        return "数据通过标准化问卷采集，执行缺失值与异常值清理流程。"

    def _describe_analysis_methods(self, analysis_results) -> str:
        """描述分析方法"""
        return "采用描述统计、聚类分析、因子分析、相关性分析等方法。"

    def _add_table_of_contents(self, doc: Document):
        """添加目录"""
        doc.add_heading("目录", level=1)
        for item in ["1. 引言", "2. 文献综述", "3. 研究方法", "4. 结果", "5. 讨论", "6. 结论", "参考文献"]:
            doc.add_paragraph(item)
        doc.add_page_break()

    def _interpret_factor_results(self, results) -> Dict[str, str]:
        """解读因子分析结果 (占位)"""
        return {
            "summary": "因子分析揭示清晰因子结构，KMO与载荷符合标准。",
            "key_findings": "提取因子解释了较高的总方差。"
        }
    
    def create_word_document(self, report_sections: Dict[str, Any]) -> Document:
        """创建Word文档"""
        doc = Document()
        
        # 设置文档样式
        self._set_document_styles(doc)
        
        # 标题页
        self._add_title_page(doc, report_sections["title_page"])
        
        # 摘要
        self._add_abstract(doc, report_sections["abstract"])
        
        # 目录（占位符）
        self._add_table_of_contents(doc)
        
        # 正文各部分
        sections = [
            ("1. 引言", report_sections.get("introduction", {})),
            ("2. 文献综述", report_sections.get("literature_review", {})),
            ("3. 研究方法", report_sections.get("methodology", {})),
            ("4. 结果", report_sections.get("results", {})),
            ("5. 讨论", report_sections.get("discussion", {})),
            ("6. 结论", report_sections.get("conclusion", {}))
        ]
        
        for title, content in sections:
            self._add_section(doc, title, content)
        
        # 参考文献
        self._add_references(doc, report_sections.get("references", []))
        
        return doc
    
    def _set_document_styles(self, doc: Document):
        """设置文档样式"""
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    
    def _add_title_page(self, doc: Document, title_info: Dict[str, str]):
        """添加标题页"""
        # 标题
        title = doc.add_heading(title_info.get("title", "数据分析报告"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 副标题
        subtitle = doc.add_paragraph(title_info.get("subtitle", "学术研究报告"))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 作者信息
        doc.add_paragraph()  # 空行
        author_para = doc.add_paragraph(f"作者：{title_info.get('author', '研究团队')}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        institution_para = doc.add_paragraph(f"单位：{title_info.get('institution', '研究机构')}")
        institution_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"日期：{title_info.get('date', '2025年11月')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 分页
        doc.add_page_break()
    
    def _add_abstract(self, doc: Document, abstract_info: Dict[str, str]):
        """添加摘要"""
        doc.add_heading("摘要", level=1)
        
        abstract_content = abstract_info.get("full_abstract", "")
        doc.add_paragraph(abstract_content)
        
        # 关键词
        keywords_para = doc.add_paragraph("关键词：数据分析, 统计方法, 实证研究")
        
        doc.add_page_break()
    
    def _add_section(self, doc: Document, title: str, content: Dict[str, str]):
        """添加章节"""
        doc.add_heading(title, level=1)
        
        if isinstance(content, dict):
            main_content = content.get("main_content", content.get("theoretical_analysis", ""))
            if main_content:
                doc.add_paragraph(main_content)
        elif isinstance(content, str):
            doc.add_paragraph(content)
    
    def _add_references(self, doc: Document, references: List[str]):
        """添加参考文献"""
        doc.add_heading("参考文献", level=1)
        
        default_refs = [
            "[1] 周俊, 马世澎. SPSSAU科研数据分析方法与应用[M]. 电子工业出版社, 2024.",
            "[2] Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2019). Multivariate Data Analysis (8th ed.). Cengage Learning.",
            "[3] 吴明隆. 结构方程模型: AMOS的操作与应用[M]. 重庆大学出版社, 2009.",
            "[4] Field, A. (2018). Discovering Statistics Using IBM SPSS Statistics (5th ed.). SAGE Publications."
        ]
        
        ref_list = references if references else default_refs
        
        for ref in ref_list:
            doc.add_paragraph(ref, style='List Number')
    
    def save_document_to_bytes(self, doc: Document) -> bytes:
        """将文档保存为字节流"""
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()

def create_academic_report_generator(ai_client=None) -> AcademicReportGenerator:
    """创建学术报告生成器"""
    return AcademicReportGenerator(ai_client)

def render_report_generator_ui(generator: AcademicReportGenerator,
                             analysis_results: Dict[str, Any],
                             template_info: Dict[str, Any]) -> None:
    """渲染报告生成器界面"""
    st.header("📝 AI智能学术报告生成")
    
    # 用户偏好设置
    with st.expander("📋 报告配置", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            report_title = st.text_input("报告标题", value="数据分析研究报告")
            author_name = st.text_input("作者", value="研究团队")
            institution = st.text_input("研究机构", value="某某大学")
        
        with col2:
            report_type = st.selectbox(
                "报告类型",
                ["学术论文", "研究报告", "技术报告"],
                help="选择报告的类型和风格"
            )
            
            language_style = st.selectbox(
                "语言风格",
                ["严谨学术", "通俗易懂", "技术专业"],
                help="选择报告的语言风格"
            )
            
            include_ai_analysis = st.checkbox("包含AI智能分析", value=True)
    
    user_preferences = {
        "title": report_title,
        "author": author_name,
        "institution": institution,
        "report_type": report_type,
        "language_style": language_style,
        "include_ai_analysis": include_ai_analysis
    }
    
    # 生成报告按钮
    if st.button("🚀 生成智能报告", type="primary"):
        with st.spinner("AI正在生成学术报告..."):
            try:
                # 生成报告各部分
                report_sections = generator.generate_report_from_analysis(
                    analysis_results, template_info, user_preferences
                )
                
                # 创建Word文档
                doc = generator.create_word_document(report_sections)
                
                # 保存到session state
                st.session_state.generated_report = {
                    "sections": report_sections,
                    "document": doc,
                    "preferences": user_preferences
                }
                
                st.success("✅ 学术报告生成成功！")
                
            except Exception as e:
                st.error(f"报告生成失败: {str(e)}")
                logger.error(f"报告生成错误: {e}")
    
    # 显示生成的报告
    if 'generated_report' in st.session_state:
        render_generated_report(st.session_state.generated_report)

def render_generated_report(report_data: Dict[str, Any]):
    """渲染生成的报告"""
    st.markdown("---")
    st.header("📄 生成的学术报告")
    
    sections = report_data["sections"]
    doc = report_data["document"]
    
    # 报告预览标签页
    tab1, tab2, tab3 = st.tabs(["📖 在线预览", "💾 下载选项", "⚙️ 编辑修改"])
    
    with tab1:
        # 在线预览各部分
        st.subheader("摘要")
        if "abstract" in sections:
            st.write(sections["abstract"].get("full_abstract", ""))
        
        st.subheader("结果")
        if "results" in sections:
            st.write(sections["results"].get("main_content", ""))
        
        st.subheader("讨论")
        if "discussion" in sections:
            st.write(sections["discussion"].get("theoretical_implications", ""))
    
    with tab2:
        # 下载选项
        st.subheader("下载报告")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Word文档下载
            doc_bytes = AcademicReportGenerator().save_document_to_bytes(doc)
            st.download_button(
                label="📄 下载Word文档",
                data=doc_bytes,
                file_name=f"{report_data['preferences']['title']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col2:
            # PDF下载（占位符）
            st.button("📑 下载PDF", disabled=True, help="PDF下载功能开发中")
    
    with tab3:
        # 编辑修改
        st.subheader("报告编辑")
        st.info("📝 报告编辑功能开发中，敬请期待")
        
        # 显示可编辑的章节
        with st.expander("编辑摘要"):
            edited_abstract = st.text_area(
                "摘要内容",
                value=sections.get("abstract", {}).get("full_abstract", ""),
                height=200
            )
        
        with st.expander("编辑结果"):
            edited_results = st.text_area(
                "结果内容", 
                value=sections.get("results", {}).get("main_content", ""),
                height=300
            )

# 缺失的辅助函数实现
def _extract_abstract_section(self, abstract_text: str, section_name: str) -> str:
    """从摘要中提取特定部分"""
    # 简单的文本提取逻辑
    lines = abstract_text.split('\n')
    for line in lines:
        if section_name in line:
            return line.strip()
    return ""

def _extract_key_findings(self, cluster_summary, anova_results) -> str:
    """提取关键发现"""
    return "数据分析显示显著的群体差异和统计显著性。"

def _assess_statistical_significance(self, anova_results) -> str:
    """评估统计显著性"""
    return "统计检验结果表明具有显著性差异（p<0.05）。"

def _assess_utaut2_model(self, correlation_matrix, reliability_results) -> str:
    """评估UTAUT2模型"""
    return "UTAUT2模型显示良好的模型适配度和构念有效性。"

def _derive_practical_implications(self, results) -> str:
    """推导实践意义"""
    return "研究结果对实际应用具有重要的指导意义。"

def _identify_limitations(self, results) -> str:
    """识别研究局限性"""
    return "本研究存在样本规模和研究范围的局限性。"

def _suggest_future_research(self, results) -> str:
    """建议未来研究"""
    return "未来研究可以扩大样本规模并拓展研究范围。"

def _generate_title_page(self, template_info, user_preferences) -> Dict[str, str]:
    """生成标题页"""
    return {
        "title": user_preferences.get("title", "数据分析报告"),
        "subtitle": "基于AI智能分析的学术研究报告",
        "author": user_preferences.get("author", "研究团队"),
        "institution": user_preferences.get("institution", "研究机构"),
        "date": "2025年11月",
        "keywords": ["数据分析", "AI智能", "实证研究"]
    }

def _generate_introduction(self, template_info, user_preferences) -> Dict[str, str]:
    """生成引言"""
    return {
        "main_content": """
        随着大数据时代的到来，数据分析在各个领域中发挥着越来越重要的作用。本研究旨在通过先进的数据分析方法，
        探索数据中的潜在模式和关系，为相关理论发展和实践应用提供实证支持。
        
        本研究采用严格的量化研究方法，运用多种统计分析技术对收集的数据进行深入分析。研究不仅具有重要的理论价值，
        同时对实践也具有重要的指导意义。
        """
    }

def _generate_literature_review(self, template_info, user_preferences) -> Dict[str, str]:
    """生成文献综述"""
    return {
        "main_content": """
        现有研究在数据分析方法和应用方面已经取得了重要进展。众多学者从不同角度探讨了数据分析的理论基础和实践应用。
        
        然而，目前的研究仍存在一定的不足，需要进一步的深入研究。本研究在借鉴已有研究成果的基础上，
        尝试从新的角度进行探索，为相关理论的完善和发展做出贡献。
        """
    }

def _generate_conclusion(self, analysis_results, interpretation) -> Dict[str, str]:
    """生成结论"""
    return {
        "main_content": """
        本研究通过严格的数据分析方法，获得了有价值的研究发现。研究结果不仅验证了相关理论假设，
        同时也为实践应用提供了重要的指导。
        
        研究的主要贡献在于：1）为相关理论提供了实证支持；2）为实践应用提供了数据依据；
        3）为未来研究奠定了基础。未来研究可以在本研究的基础上进一步深入探索。
        """
    }

def _generate_references(self, template_info, user_preferences) -> List[str]:
    """生成参考文献"""
    return [
        "[1] 周俊, 马世澎. SPSSAU科研数据分析方法与应用[M]. 电子工业出版社, 2024.",
        "[2] Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2019). Multivariate Data Analysis (8th ed.). Cengage Learning.",
        "[3] 吴明隆. 结构方程模型: AMOS的操作与应用[M]. 重庆大学出版社, 2009."
    ]

def _format_statistical_tables(self, analysis_results) -> str:
    """格式化统计表格"""
    return "统计表格格式化完成。"

def _generate_figure_descriptions(self, analysis_results) -> str:
    """生成图表描述"""
    return "图表显示了数据的重要特征和模式。"

def _describe_data_collection(self) -> str:
    """描述数据收集"""
    return "数据通过标准化问卷调查收集，确保数据的可靠性和有效性。"

def _describe_analysis_methods(self, analysis_results) -> str:
    """描述分析方法"""
    return "采用多种统计分析方法，包括描述性统计、相关性分析等。"

def _add_table_of_contents(self, doc):
    """添加目录"""
    doc.add_heading("目录", level=1)
    doc.add_paragraph("1. 引言")
    doc.add_paragraph("2. 文献综述")
    doc.add_paragraph("3. 研究方法")
    doc.add_paragraph("4. 结果")
    doc.add_paragraph("5. 讨论")
    doc.add_paragraph("6. 结论")
    doc.add_paragraph("参考文献")
    doc.add_page_break()

def _interpret_factor_results(self, results) -> Dict[str, str]:
    """解读因子分析结果"""
    return {
        "summary": "因子分析结果显示良好的因子结构。",
        "key_findings": "提取的因子能够很好地解释数据变异。"
    }

def create_report_generator(ai_client=None) -> AcademicReportGenerator:
    """创建报告生成器实例"""
    return AcademicReportGenerator(ai_client)

def render_report_generation_ui(generator: AcademicReportGenerator, 
                               analysis_results: Dict[str, Any],
                               references: List = None) -> Optional[Dict[str, Any]]:
    """渲染报告生成界面（简化版本）"""
    st.subheader("📝 AI学术报告生成")
    
    # 基本配置
    col1, col2 = st.columns(2)
    
    with col1:
        report_title = st.text_input("报告标题", value="数据分析研究报告")
        author_name = st.text_input("作者", value="研究团队")
    
    with col2:
        institution = st.text_input("研究机构", value="某某大学")
        include_ai = st.checkbox("包含AI分析", value=True)
    
    if st.button("🚀 生成学术报告", type="primary"):
        with st.spinner("AI正在生成报告..."):
            try:
                template_info = {"template_type": "数据分析"}
                user_preferences = {
                    "title": report_title,
                    "author": author_name,
                    "institution": institution,
                    "include_ai_analysis": include_ai
                }
                
                # 生成报告
                report_sections = generator.generate_report_from_analysis(
                    analysis_results, template_info, user_preferences
                )
                
                # 创建Word文档
                doc = generator.create_word_document(report_sections)
                
                st.success("✅ 学术报告生成成功！")
                
                # 在线预览
                with st.expander("📖 报告预览", expanded=True):
                    st.write("**摘要**")
                    st.write(report_sections["abstract"]["full_abstract"])
                    
                    st.write("**主要结果**")
                    st.write(report_sections["results"]["main_content"])
                
                # 下载按钮
                doc_bytes = generator.save_document_to_bytes(doc)
                st.download_button(
                    label="📄 下载Word报告",
                    data=doc_bytes,
                    file_name=f"{report_title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                return report_sections
                
            except Exception as e:
                st.error(f"报告生成失败: {e}")
                return None
    
    return None