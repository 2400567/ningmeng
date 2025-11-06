"""
用户报告样例上传和模板管理模块
允许用户上传参考报告，系统自动学习和应用模板风格
"""

import os
import json
import hashlib
from typing import Dict, List, Optional, Any
from datetime import datetime
import streamlit as st
import pandas as pd
from docx import Document
import PyPDF2
import re

class ReportTemplateUploader:
    """报告模板上传器"""
    
    def __init__(self):
        self.upload_dir = "temp/uploaded_templates"
        self.template_db_path = "temp/template_database.json"
        self.ensure_directories()
        self.template_analyzer = TemplateAnalyzer()
        
    def ensure_directories(self):
        """确保必要目录存在"""
        os.makedirs(self.upload_dir, exist_ok=True)
        os.makedirs("temp", exist_ok=True)
    
    def upload_template(self, uploaded_file, template_info: Dict) -> Dict[str, Any]:
        """
        上传报告模板
        
        Args:
            uploaded_file: Streamlit上传的文件对象
            template_info: 模板信息
            
        Returns:
            上传结果
        """
        try:
            # 生成文件哈希
            file_content = uploaded_file.getvalue()
            file_hash = hashlib.md5(file_content).hexdigest()
            
            # 保存文件
            file_extension = uploaded_file.name.split('.')[-1].lower()
            saved_filename = f"{file_hash}.{file_extension}"
            saved_path = os.path.join(self.upload_dir, saved_filename)
            
            with open(saved_path, "wb") as f:
                f.write(file_content)
            
            # 分析模板内容
            analysis_result = self.template_analyzer.analyze_template(saved_path, file_extension)
            
            # 保存模板信息到数据库
            template_record = {
                "id": file_hash,
                "original_name": uploaded_file.name,
                "saved_path": saved_path,
                "file_type": file_extension,
                "upload_time": datetime.now().isoformat(),
                "template_info": template_info,
                "analysis_result": analysis_result,
                "usage_count": 0
            }
            
            self._save_template_to_db(template_record)
            
            return {
                "success": True,
                "template_id": file_hash,
                "message": "模板上传成功",
                "analysis": analysis_result
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"上传失败: {str(e)}"
            }
    
    def get_available_templates(self) -> List[Dict]:
        """获取可用模板列表"""
        try:
            if os.path.exists(self.template_db_path):
                with open(self.template_db_path, 'r', encoding='utf-8') as f:
                    db = json.load(f)
                    return list(db.values())
            return []
        except Exception:
            return []
    
    def get_template_by_id(self, template_id: str) -> Optional[Dict]:
        """根据ID获取模板"""
        templates = self.get_available_templates()
        for template in templates:
            if template["id"] == template_id:
                return template
        return None
    
    def delete_template(self, template_id: str) -> bool:
        """删除模板"""
        try:
            template = self.get_template_by_id(template_id)
            if template:
                # 删除文件
                if os.path.exists(template["saved_path"]):
                    os.remove(template["saved_path"])
                
                # 从数据库删除记录
                self._remove_template_from_db(template_id)
                return True
            return False
        except Exception:
            return False
    
    def _save_template_to_db(self, template_record: Dict):
        """保存模板到数据库"""
        db = {}
        if os.path.exists(self.template_db_path):
            with open(self.template_db_path, 'r', encoding='utf-8') as f:
                try:
                    db = json.load(f)
                except json.JSONDecodeError:
                    db = {}
        
        db[template_record["id"]] = template_record
        
        with open(self.template_db_path, 'w', encoding='utf-8') as f:
            json.dump(db, f, ensure_ascii=False, indent=2)
    
    def _remove_template_from_db(self, template_id: str):
        """从数据库删除模板"""
        if os.path.exists(self.template_db_path):
            with open(self.template_db_path, 'r', encoding='utf-8') as f:
                try:
                    db = json.load(f)
                    if template_id in db:
                        del db[template_id]
                        
                        with open(self.template_db_path, 'w', encoding='utf-8') as f:
                            json.dump(db, f, ensure_ascii=False, indent=2)
                except json.JSONDecodeError:
                    pass

class TemplateAnalyzer:
    """模板分析器"""
    
    def analyze_template(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        分析模板文件结构和风格
        
        Args:
            file_path: 文件路径
            file_type: 文件类型
            
        Returns:
            分析结果
        """
        if file_type in ['docx', 'doc']:
            return self._analyze_word_template(file_path)
        elif file_type == 'pdf':
            return self._analyze_pdf_template(file_path)
        elif file_type in ['txt', 'md']:
            return self._analyze_text_template(file_path)
        else:
            return {"error": "不支持的文件类型"}
    
    def _analyze_word_template(self, file_path: str) -> Dict[str, Any]:
        """分析Word文档模板"""
        try:
            doc = Document(file_path)
            
            analysis = {
                "structure": [],
                "sections": {},
                "style_info": {},
                "content_patterns": [],
                "statistics": {}
            }
            
            # 分析文档结构
            current_section = None
            section_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                # 检测标题级别
                if self._is_heading(paragraph):
                    if current_section:
                        analysis["sections"][current_section] = {
                            "content": section_content,
                            "word_count": sum(len(content.split()) for content in section_content)
                        }
                    
                    current_section = text
                    section_content = []
                    analysis["structure"].append({
                        "type": "heading",
                        "text": text,
                        "level": self._get_heading_level(paragraph)
                    })
                else:
                    section_content.append(text)
                    analysis["structure"].append({
                        "type": "paragraph",
                        "text": text,
                        "section": current_section
                    })
            
            # 保存最后一个章节
            if current_section:
                analysis["sections"][current_section] = {
                    "content": section_content,
                    "word_count": sum(len(content.split()) for content in section_content)
                }
            
            # 提取内容模式
            analysis["content_patterns"] = self._extract_content_patterns(analysis["sections"])
            
            # 统计信息
            analysis["statistics"] = {
                "total_paragraphs": len(doc.paragraphs),
                "total_sections": len(analysis["sections"]),
                "total_words": sum(len(p.text.split()) for p in doc.paragraphs),
                "average_section_length": sum(s["word_count"] for s in analysis["sections"].values()) / len(analysis["sections"]) if analysis["sections"] else 0
            }
            
            return analysis
            
        except Exception as e:
            return {"error": f"Word文档分析失败: {str(e)}"}
    
    def _analyze_pdf_template(self, file_path: str) -> Dict[str, Any]:
        """分析PDF模板"""
        try:
            analysis = {
                "structure": [],
                "sections": {},
                "content_patterns": [],
                "statistics": {}
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    full_text += page_text + "\n"
                
                # 分析文本结构
                lines = full_text.split('\n')
                current_section = None
                section_content = []
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # 简单的标题检测
                    if self._is_likely_heading(line):
                        if current_section:
                            analysis["sections"][current_section] = {
                                "content": section_content,
                                "word_count": sum(len(content.split()) for content in section_content)
                            }
                        
                        current_section = line
                        section_content = []
                        analysis["structure"].append({
                            "type": "heading",
                            "text": line
                        })
                    else:
                        section_content.append(line)
                        analysis["structure"].append({
                            "type": "paragraph",
                            "text": line,
                            "section": current_section
                        })
                
                # 保存最后一个章节
                if current_section:
                    analysis["sections"][current_section] = {
                        "content": section_content,
                        "word_count": sum(len(content.split()) for content in section_content)
                    }
                
                # 统计信息
                analysis["statistics"] = {
                    "total_pages": len(pdf_reader.pages),
                    "total_sections": len(analysis["sections"]),
                    "total_words": len(full_text.split()),
                    "average_section_length": sum(s["word_count"] for s in analysis["sections"].values()) / len(analysis["sections"]) if analysis["sections"] else 0
                }
            
            return analysis
            
        except Exception as e:
            return {"error": f"PDF分析失败: {str(e)}"}
    
    def _analyze_text_template(self, file_path: str) -> Dict[str, Any]:
        """分析文本模板"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            analysis = {
                "structure": [],
                "sections": {},
                "content_patterns": [],
                "statistics": {}
            }
            
            lines = content.split('\n')
            current_section = None
            section_content = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Markdown标题检测
                if line.startswith('#'):
                    if current_section:
                        analysis["sections"][current_section] = {
                            "content": section_content,
                            "word_count": sum(len(content.split()) for content in section_content)
                        }
                    
                    current_section = line.lstrip('#').strip()
                    section_content = []
                    analysis["structure"].append({
                        "type": "heading",
                        "text": current_section,
                        "level": len(line) - len(line.lstrip('#'))
                    })
                else:
                    section_content.append(line)
                    analysis["structure"].append({
                        "type": "paragraph",
                        "text": line,
                        "section": current_section
                    })
            
            # 保存最后一个章节
            if current_section:
                analysis["sections"][current_section] = {
                    "content": section_content,
                    "word_count": sum(len(content.split()) for content in section_content)
                }
            
            # 统计信息
            analysis["statistics"] = {
                "total_lines": len(lines),
                "total_sections": len(analysis["sections"]),
                "total_words": len(content.split()),
                "average_section_length": sum(s["word_count"] for s in analysis["sections"].values()) / len(analysis["sections"]) if analysis["sections"] else 0
            }
            
            return analysis
            
        except Exception as e:
            return {"error": f"文本分析失败: {str(e)}"}
    
    def _is_heading(self, paragraph) -> bool:
        """判断段落是否为标题"""
        try:
            # 检查样式名称
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name or 'title' in style_name:
                return True
            
            # 检查字体大小和格式
            if paragraph.runs:
                run = paragraph.runs[0]
                if run.font.size and run.font.size.pt > 14:
                    return True
                if run.bold:
                    return True
            
            return False
        except:
            return False
    
    def _get_heading_level(self, paragraph) -> int:
        """获取标题级别"""
        try:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name:
                # 提取数字
                numbers = re.findall(r'\d+', style_name)
                if numbers:
                    return int(numbers[0])
            return 1
        except:
            return 1
    
    def _is_likely_heading(self, line: str) -> bool:
        """判断行是否可能是标题"""
        # 简单的启发式规则
        if len(line) < 100 and len(line) > 3:
            # 检查是否包含常见标题关键词
            heading_keywords = [
                '摘要', '引言', '背景', '方法', '结果', '讨论', '结论',
                '参考文献', '致谢', '附录', '概述', '分析', '总结',
                'Abstract', 'Introduction', 'Methods', 'Results', 
                'Discussion', 'Conclusion', 'References'
            ]
            
            for keyword in heading_keywords:
                if keyword in line:
                    return True
            
            # 检查是否为数字编号
            if re.match(r'^\d+[\.\s]', line):
                return True
        
        return False
    
    def _extract_content_patterns(self, sections: Dict) -> List[Dict]:
        """提取内容模式"""
        patterns = []
        
        for section_name, section_data in sections.items():
            content = " ".join(section_data["content"])
            
            # 分析内容类型
            pattern = {
                "section": section_name,
                "type": self._classify_content_type(content),
                "length": section_data["word_count"],
                "common_phrases": self._extract_common_phrases(content)
            }
            
            patterns.append(pattern)
        
        return patterns
    
    def _classify_content_type(self, content: str) -> str:
        """分类内容类型"""
        content_lower = content.lower()
        
        if any(keyword in content_lower for keyword in ['表', '图', '数据', '统计', '分析']):
            return "数据分析"
        elif any(keyword in content_lower for keyword in ['方法', '步骤', '过程', '流程']):
            return "方法描述"
        elif any(keyword in content_lower for keyword in ['结果', '发现', '显示', '表明']):
            return "结果报告"
        elif any(keyword in content_lower for keyword in ['讨论', '意义', '影响', '解释']):
            return "讨论分析"
        elif any(keyword in content_lower for keyword in ['结论', '总结', '建议', '展望']):
            return "结论建议"
        else:
            return "一般文本"
    
    def _extract_common_phrases(self, content: str) -> List[str]:
        """提取常见短语"""
        # 简单的短语提取
        sentences = re.split(r'[。！？]', content)
        phrases = []
        
        for sentence in sentences[:5]:  # 只取前5个句子
            sentence = sentence.strip()
            if len(sentence) > 10 and len(sentence) < 100:
                phrases.append(sentence)
        
        return phrases

class TemplateApplier:
    """模板应用器"""
    
    def __init__(self):
        self.uploader = ReportTemplateUploader()
    
    def apply_template_style(self, 
                           content: Dict, 
                           template_id: str, 
                           report_type: str = "academic") -> Dict[str, Any]:
        """
        应用模板风格到报告内容
        
        Args:
            content: 报告内容
            template_id: 模板ID
            report_type: 报告类型
            
        Returns:
            应用模板后的报告
        """
        template = self.uploader.get_template_by_id(template_id)
        if not template:
            return {"error": "模板不存在"}
        
        try:
            analysis = template["analysis_result"]
            
            # 应用结构模式
            structured_content = self._apply_structure_pattern(content, analysis)
            
            # 应用内容风格
            styled_content = self._apply_content_style(structured_content, analysis)
            
            # 更新使用次数
            self._update_template_usage(template_id)
            
            return {
                "success": True,
                "content": styled_content,
                "template_applied": template["original_name"],
                "style_notes": self._generate_style_notes(analysis)
            }
            
        except Exception as e:
            return {"error": f"模板应用失败: {str(e)}"}
    
    def _apply_structure_pattern(self, content: Dict, analysis: Dict) -> Dict:
        """应用结构模式"""
        if "structure" not in analysis:
            return content
        
        # 重新组织内容结构
        structured = content.copy()
        
        # 根据模板结构调整章节顺序
        template_sections = [item["text"] for item in analysis["structure"] if item["type"] == "heading"]
        
        if template_sections and "sections" in structured:
            # 尝试匹配和重新排序章节
            reordered_sections = {}
            
            for template_section in template_sections:
                # 寻找最匹配的章节
                best_match = self._find_best_section_match(template_section, structured["sections"])
                if best_match:
                    reordered_sections[template_section] = structured["sections"][best_match]
            
            # 添加未匹配的章节
            for section_name, section_content in structured["sections"].items():
                if not any(section_name in reordered_sections.values() for reordered_sections in [reordered_sections]):
                    reordered_sections[section_name] = section_content
            
            structured["sections"] = reordered_sections
        
        return structured
    
    def _apply_content_style(self, content: Dict, analysis: Dict) -> Dict:
        """应用内容风格"""
        styled = content.copy()
        
        if "content_patterns" in analysis:
            for section_name, section_content in styled.get("sections", {}).items():
                # 寻找匹配的内容模式
                pattern = self._find_matching_pattern(section_name, analysis["content_patterns"])
                
                if pattern and "common_phrases" in pattern:
                    # 应用常见短语风格
                    if isinstance(section_content, str):
                        styled_content = self._apply_phrase_style(section_content, pattern["common_phrases"])
                        styled["sections"][section_name] = styled_content
        
        return styled
    
    def _find_best_section_match(self, template_section: str, content_sections: Dict) -> Optional[str]:
        """寻找最佳章节匹配"""
        template_keywords = set(template_section.lower().split())
        
        best_match = None
        best_score = 0
        
        for section_name in content_sections.keys():
            section_keywords = set(section_name.lower().split())
            
            # 计算词汇重叠度
            overlap = len(template_keywords & section_keywords)
            score = overlap / len(template_keywords) if template_keywords else 0
            
            if score > best_score and score > 0.3:  # 至少30%匹配度
                best_score = score
                best_match = section_name
        
        return best_match
    
    def _find_matching_pattern(self, section_name: str, patterns: List[Dict]) -> Optional[Dict]:
        """寻找匹配的内容模式"""
        for pattern in patterns:
            if self._sections_similar(section_name, pattern["section"]):
                return pattern
        return None
    
    def _sections_similar(self, section1: str, section2: str) -> bool:
        """判断章节是否相似"""
        keywords1 = set(section1.lower().split())
        keywords2 = set(section2.lower().split())
        
        if not keywords1 or not keywords2:
            return False
        
        overlap = len(keywords1 & keywords2)
        similarity = overlap / min(len(keywords1), len(keywords2))
        
        return similarity > 0.5
    
    def _apply_phrase_style(self, content: str, common_phrases: List[str]) -> str:
        """应用短语风格"""
        # 简单的风格应用 - 在实际应用中可以更复杂
        styled_content = content
        
        # 这里可以实现更复杂的风格转换逻辑
        # 例如：句式调整、用词替换等
        
        return styled_content
    
    def _update_template_usage(self, template_id: str):
        """更新模板使用次数"""
        try:
            if os.path.exists(self.uploader.template_db_path):
                with open(self.uploader.template_db_path, 'r', encoding='utf-8') as f:
                    db = json.load(f)
                
                if template_id in db:
                    db[template_id]["usage_count"] += 1
                    db[template_id]["last_used"] = datetime.now().isoformat()
                    
                    with open(self.uploader.template_db_path, 'w', encoding='utf-8') as f:
                        json.dump(db, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
    
    def _generate_style_notes(self, analysis: Dict) -> List[str]:
        """生成风格说明"""
        notes = []
        
        if "statistics" in analysis:
            stats = analysis["statistics"]
            notes.append(f"模板包含 {stats.get('total_sections', 0)} 个主要章节")
            if "average_section_length" in stats:
                notes.append(f"平均章节长度: {stats['average_section_length']:.0f} 字")
        
        if "content_patterns" in analysis:
            content_types = [pattern["type"] for pattern in analysis["content_patterns"]]
            unique_types = list(set(content_types))
            notes.append(f"内容类型: {', '.join(unique_types)}")
        
        return notes

# Streamlit界面组件
def create_template_upload_interface():
    """创建模板上传界面"""
    st.subheader("📄 上传报告模板")
    
    uploader = ReportTemplateUploader()
    
    # 上传文件
    uploaded_file = st.file_uploader(
        "选择报告模板文件",
        type=['docx', 'doc', 'pdf', 'txt', 'md'],
        help="支持Word文档、PDF、文本文件和Markdown文件"
    )
    
    if uploaded_file is not None:
        # 模板信息输入
        col1, col2 = st.columns(2)
        
        with col1:
            template_name = st.text_input("模板名称", value=uploaded_file.name.split('.')[0])
            template_type = st.selectbox(
                "模板类型",
                ["学术论文", "商业报告", "技术报告", "研究提案", "其他"]
            )
        
        with col2:
            template_category = st.selectbox(
                "适用领域",
                ["通用", "经济管理", "工程技术", "社会科学", "自然科学", "医学", "其他"]
            )
            template_language = st.selectbox("语言", ["中文", "英文", "中英混合"])
        
        template_description = st.text_area("模板描述", placeholder="描述这个模板的特点和用途...")
        
        if st.button("上传并分析模板", type="primary"):
            template_info = {
                "name": template_name,
                "type": template_type,
                "category": template_category,
                "language": template_language,
                "description": template_description
            }
            
            with st.spinner("正在上传和分析模板..."):
                result = uploader.upload_template(uploaded_file, template_info)
            
            if result["success"]:
                st.success(result["message"])
                st.info(f"模板ID: {result['template_id']}")
                
                # 显示分析结果
                if "analysis" in result and "error" not in result["analysis"]:
                    st.subheader("📊 模板分析结果")
                    analysis = result["analysis"]
                    
                    if "statistics" in analysis:
                        stats = analysis["statistics"]
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            st.metric("章节数量", stats.get("total_sections", 0))
                        with col2:
                            st.metric("总字数", stats.get("total_words", 0))
                        with col3:
                            if "average_section_length" in stats:
                                st.metric("平均章节长度", f"{stats['average_section_length']:.0f}")
                    
                    # 显示章节结构
                    if "sections" in analysis and analysis["sections"]:
                        st.subheader("📋 章节结构")
                        for section_name, section_info in analysis["sections"].items():
                            with st.expander(f"{section_name} ({section_info['word_count']} 字)"):
                                if section_info["content"]:
                                    st.write("内容预览:")
                                    preview = " ".join(section_info["content"][:2])  # 显示前两段
                                    st.text(preview[:200] + "..." if len(preview) > 200 else preview)
            else:
                st.error(result["message"])

def create_template_management_interface():
    """创建模板管理界面"""
    st.subheader("📚 模板管理")
    
    uploader = ReportTemplateUploader()
    templates = uploader.get_available_templates()
    
    if not templates:
        st.info("暂无已上传的模板")
        return
    
    # 显示模板列表
    for template in templates:
        with st.expander(f"📄 {template['template_info']['name']} ({template['template_info']['type']})"):
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col1:
                st.write(f"**描述:** {template['template_info'].get('description', '无描述')}")
                st.write(f"**类型:** {template['template_info']['type']}")
                st.write(f"**领域:** {template['template_info']['category']}")
                st.write(f"**语言:** {template['template_info']['language']}")
                st.write(f"**上传时间:** {template['upload_time'][:19]}")
                st.write(f"**使用次数:** {template.get('usage_count', 0)}")
            
            with col2:
                if "analysis_result" in template and "statistics" in template["analysis_result"]:
                    stats = template["analysis_result"]["statistics"]
                    st.metric("章节数", stats.get("total_sections", 0))
                    st.metric("字数", stats.get("total_words", 0))
            
            with col3:
                if st.button(f"删除", key=f"delete_{template['id']}"):
                    if uploader.delete_template(template["id"]):
                        st.success("模板已删除")
                        st.rerun()
                    else:
                        st.error("删除失败")

def create_template_selection_interface():
    """创建模板选择界面"""
    uploader = ReportTemplateUploader()
    templates = uploader.get_available_templates()
    
    if not templates:
        st.info("请先上传报告模板")
        return None
    
    # 模板选择
    template_options = {
        f"{t['template_info']['name']} ({t['template_info']['type']})": t['id'] 
        for t in templates
    }
    
    selected_template_name = st.selectbox("选择报告模板", options=list(template_options.keys()))
    
    if selected_template_name:
        template_id = template_options[selected_template_name]
        template = uploader.get_template_by_id(template_id)
        
        # 显示模板信息
        with st.expander("模板详情"):
            st.write(f"**类型:** {template['template_info']['type']}")
            st.write(f"**领域:** {template['template_info']['category']}")
            st.write(f"**描述:** {template['template_info'].get('description', '无描述')}")
            
            if "analysis_result" in template and "statistics" in template["analysis_result"]:
                stats = template["analysis_result"]["statistics"]
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("章节数", stats.get("total_sections", 0))
                with col2:
                    st.metric("总字数", stats.get("total_words", 0))
                with col3:
                    st.metric("使用次数", template.get("usage_count", 0))
        
        return template_id
    
    return None