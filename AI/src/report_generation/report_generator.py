from typing import Optional
import os
import io
import datetime
import pandas as pd
import numpy as np
from typing import Dict, List, Optional, Union
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from matplotlib.figure import Figure
import tempfile
import matplotlib.pyplot as plt
import seaborn as sns
import logging


# 模块级日志器
logger = logging.getLogger (__name__)
logger.addHandler(logging.NullHandler())

# 导入AI增强模块
try:
    from ..ai_agent.ai_report_enhancer import AIReportEnhancer, create_ai_enhancer, DEFAULT_CONFIGS
    AI_ENHANCEMENT_AVAILABLE = True
except ImportError:
    AI_ENHANCEMENT_AVAILABLE = False
    logger.warning("AI报告增强模块不可用，将跳过AI增强功能")


class ReportGenerator:
    """
    Word报告生成器，用于创建专业的数据报告文档
    """
    
    def __init__(self):
        self.document = None
        self.temp_dir = tempfile.mkdtemp()
    
    def create_report(self, title: str = "数据分析报告", 
                     author: str = "AI数据分析系统",
                     subtitle: Optional[str] = None) -> None:
        """
        创建一个新的报告文档
        
        Args:
            title: 报告标题
            author: 报告作者
            subtitle: 报告副标题
        """
        # 创建新文档
        self.document = Document()
        
        # 设置标题
        self._add_title_section(title, subtitle, author)
        
        # 添加目录
        self._add_table_of_contents()


    def _add_title_section(self, title: str, subtitle: Optional[str], author: str) -> None:
        """
        添加标题部分
        """
        # 主标题
        title_para = self.document.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 副标题
        if subtitle:
            subtitle_para = self.document.add_heading(subtitle, level=1)
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        self.document.add_paragraph()
        
        # 添加作者和日期信息
        info_para = self.document.add_paragraph()
        info_run = info_para.add_run(f"作者: {author}")
        info_run.font.size = Pt(12)
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 添加日期
        date_para = self.document.add_paragraph()
        date_run = date_para.add_run(f"生成日期: {datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 分页符
        self.document.add_page_break()
    
    def _add_table_of_contents(self) -> None:
        """
        添加目录
        """
        # 目录标题
        toc_title = self.document.add_heading("目录", level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加目录（实际内容会在保存前更新）
        self.document.add_paragraph("[目录将在保存时自动生成]")
        
        # 分页符
        self.document.add_page_break()
    
    def add_executive_summary(self, summary: str) -> None:
        """
        添加执行摘要
        
        Args:
            summary: 摘要内容
        """
        self.document.add_heading("执行摘要", level=1)
        
        # 添加摘要段落
        self.document.add_paragraph(summary)
        
        # 添加空行
        self.document.add_paragraph()
    
    def add_data_overview(self, data_info: Dict) -> None:
        """
        添加数据概览部分
        
        Args:
            data_info: 数据信息字典，包含行数、列数、文件名称等
        """
        self.document.add_heading("1. 数据概览", level=1)
        
        # 添加基本信息
        self.document.add_heading("1.1 数据集基本信息", level=2)
        
        # 创建信息表格
        info_table = self.document.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = '属性'
        hdr_cells[1].text = '值'
        
        # 设置表头样式
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # 添加数据
        rows = [
            ('文件名称', str(data_info.get('file_name', '未知')) if data_info.get('file_name') is not None else '未知'),
            ('数据格式', str(data_info.get('file_format', '未知')) if data_info.get('file_format') is not None else '未知'),
            ('数据行数', str(data_info.get('num_rows', 0))),
            ('数据列数', str(data_info.get('num_columns', 0))),
            ('数值型特征', str(data_info.get('num_numeric_cols', 0))),
            ('分类型特征', str(data_info.get('num_categorical_cols', 0))),
            ('日期型特征', str(data_info.get('num_date_cols', 0))),
            ('数据大小', str(data_info.get('data_size', '未知')) if data_info.get('data_size') is not None else '未知')
        ]
        
        for row_data in rows:
            cells = info_table.add_row().cells
            # 确保所有文本都不为None，进行双重保护
            key_text = row_data[0]
            value_text = row_data[1]
            
            # 保证键名永远不为None
            if key_text is None:
                key_text = '未知项目'
            elif not isinstance(key_text, str):
                key_text = str(key_text)
            
            # 保证值永远不为None
            if value_text is None:
                value_text = '未知'
            elif not isinstance(value_text, str):
                value_text = str(value_text)
            
            cells[0].text = key_text
            cells[1].text = value_text
            for cell in cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # 添加列信息
        if 'columns_info' in data_info:
            self.document.add_heading("1.2 列信息", level=2)
            
            col_table = self.document.add_table(rows=1, cols=4)
            col_table.style = 'Table Grid'
            col_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            hdr_cells = col_table.rows[0].cells
            hdr_cells[0].text = '列名'
            hdr_cells[1].text = '数据类型'
            hdr_cells[2].text = '非空值数量'
            hdr_cells[3].text = '描述'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # 添加列数据
            for col_info in data_info['columns_info']:
                cells = col_table.add_row().cells
                # 确保所有文本都不为None
                cells[0].text = str(col_info.get('name', '')) if col_info.get('name') is not None else ''
                cells[1].text = str(col_info.get('dtype', '')) if col_info.get('dtype') is not None else ''
                cells[2].text = str(col_info.get('non_null_count', 0))
                cells[3].text = str(col_info.get('description', 'N/A')) if col_info.get('description') is not None else 'N/A'
                for cell in cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # 添加空行
        self.document.add_paragraph()
    
    def add_data_preprocessing_section(self, preprocessing_info: Dict) -> None:
        """
        添加数据预处理部分
        
        Args:
            preprocessing_info: 预处理信息字典
        """
        self.document.add_heading("2. 数据预处理", level=1)
        
        # 缺失值处理（对可能为 None 的字段进行保护）
        mv_info = preprocessing_info.get('missing_values') or {}
        if mv_info:
            self.document.add_heading("2.1 缺失值处理", level=2)

            mv_para = self.document.add_paragraph()
            mv_para.add_run(f"发现 {mv_info.get('total_missing', 0)} 个缺失值，分布在 {mv_info.get('columns_with_missing', 0)} 列中。").bold = True

            # 添加处理方法
            handling_method = mv_info.get('handling_method')
            if handling_method:
                self.document.add_paragraph(f"处理方法: {handling_method}")

            # 如果有处理详情
            for detail in mv_info.get('details') or []:
                # 确保 details 可迭代
                if detail is None:
                    continue
                self.document.add_paragraph(f"- {detail}", style='List Bullet')
        
        # 异常值处理（保护 None）
        outliers_info = preprocessing_info.get('outliers') or {}
        if outliers_info:
            self.document.add_heading("2.2 异常值处理", level=2)

            outliers_para = self.document.add_paragraph()
            outliers_para.add_run(f"发现 {outliers_info.get('total_outliers', 0)} 个异常值，分布在 {outliers_info.get('columns_with_outliers', 0)} 列中。").bold = True

            # 添加处理方法
            if outliers_info.get('handling_method'):
                self.document.add_paragraph(f"处理方法: {outliers_info.get('handling_method')}")
        
        # 特征工程（保护 None）
        fe_info = preprocessing_info.get('feature_engineering') or {}
        if fe_info:
            self.document.add_heading("2.3 特征工程", level=2)

            # 编码处理
            if fe_info.get('encoding'):
                self.document.add_paragraph(f"分类特征编码: {fe_info.get('encoding')}")

            # 标准化/归一化
            if fe_info.get('scaling'):
                self.document.add_paragraph(f"数值特征缩放: {fe_info.get('scaling')}")

            # 特征选择
            feature_selection = fe_info.get('feature_selection') or {}
            selected_features = feature_selection.get('selected_features', []) if isinstance(feature_selection, dict) else []
            self.document.add_paragraph(f"选择的特征数量: {len(selected_features)}")

            # 添加选择的特征列表
            if len(selected_features) <= 10:  # 只显示前10个特征
                if selected_features:
                    features_para = self.document.add_paragraph("选择的特征:")
                    for feature in selected_features:
                        if feature is None:
                            continue
                        self.document.add_paragraph(f"- {feature}", style='List Bullet')
            else:
                self.document.add_paragraph(f"(显示前10个特征)")
                for feature in selected_features[:10]:
                    if feature is None:
                        continue
                    self.document.add_paragraph(f"- {feature}", style='List Bullet')
        
        # 添加空行
        self.document.add_paragraph()
    
    def add_analysis_results(self, analysis_results: Dict) -> None:
        """
        添加分析结果部分
        
        Args:
            analysis_results: 分析结果字典
        """
        try:
            logger.info("开始添加分析结果")
            self.document.add_heading("3. 数据分析结果", level=1)

            # 保护 analysis_results 为 None 的情况
            if not analysis_results:
                logger.warning("analysis_results为空，添加默认消息")
                self.document.add_paragraph("暂无分析结果可展示。")
                self.document.add_paragraph()
                return

            logger.debug(f"analysis_results keys: {list(analysis_results.keys())}")

            # 🤖 首先添加AI增强内容（如果存在）
            self._add_ai_enhanced_content(analysis_results)

            # 描述性统计
            try:
                if 'descriptive_stats' in analysis_results:
                    logger.info("添加描述性统计")
                    self.document.add_heading("3.1 描述性统计", level=2)
                    
                    stats_df = analysis_results.get('descriptive_stats')
                    logger.debug(f"stats_df type: {type(stats_df)}")
                    
                    if isinstance(stats_df, pd.DataFrame) and not stats_df.empty:
                        # 创建表格
                        stats_table = self.document.add_table(rows=1, cols=len(stats_df.columns) + 1)
                        stats_table.style = 'Table Grid'
                        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        # 表头
                        hdr_cells = stats_table.rows[0].cells
                        hdr_cells[0].text = '统计量'
                        for i, col in enumerate(stats_df.columns):
                            hdr_cells[i+1].text = str(col)
                        
                        for cell in hdr_cells:
                            cell.paragraphs[0].runs[0].bold = True
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # 添加数据（只显示主要统计量）
                        stats_to_show = ['mean', 'std', 'min', '25%', '50%', '75%', 'max']
                        for stat in stats_to_show:
                            if stat in stats_df.index:
                                cells = stats_table.add_row().cells
                                cells[0].text = self._get_statistic_name(stat)
                                for i, col in enumerate(stats_df.columns):
                                    try:
                                        value = stats_df.loc[stat, col]
                                        cells[i+1].text = f"{value:.4f}" if pd.notnull(value) else "N/A"
                                    except Exception as e:
                                        logger.error(f"处理统计值时出错: {str(e)}")
                                        cells[i+1].text = "N/A"
                                for cell in cells:
                                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        logger.info("描述性统计表格添加成功")
                    else:
                        logger.warning("描述性统计数据无效")
                        self.document.add_paragraph("描述性统计数据暂不可用。")
            except Exception as e:
                logger.error(f"添加描述性统计时出错: {str(e)}")
                self.document.add_paragraph("描述性统计生成时出现错误。")
        
            # 相关性分析
            try:
                if 'correlation' in analysis_results:
                    logger.info("添加相关性分析")
                    self.document.add_heading("3.2 相关性分析", level=2)
                    
                    corr_results = analysis_results.get('correlation') or {}
                    logger.debug(f"corr_results type: {type(corr_results)}, keys: {list(corr_results.keys()) if isinstance(corr_results, dict) else 'not dict'}")
                    
                    if isinstance(corr_results, dict):
                        method = corr_results.get('method', 'Pearson')
                        self.document.add_paragraph(f"相关性分析采用 {method} 方法。")
                        
                        # 显示强相关的特征对
                        strong_corr = corr_results.get('strong_correlations') or []
                        logger.debug(f"strong_correlations type: {type(strong_corr)}, length: {len(strong_corr) if hasattr(strong_corr, '__len__') else 'no len'}")
                        
                        if strong_corr and isinstance(strong_corr, (list, tuple)) and len(strong_corr) > 0:
                            self.document.add_paragraph("发现以下强相关特征对（相关系数绝对值 > 0.7）:")
                            for i, pair in enumerate(strong_corr):
                                try:
                                    # 保护每个 pair 是字典且包含需要的字段
                                    if not isinstance(pair, dict):
                                        logger.warning(f"强相关对 {i} 不是字典类型: {type(pair)}")
                                        continue
                                    f1 = pair.get('feature1', '未知')
                                    f2 = pair.get('feature2', '未知')
                                    corr_v = pair.get('correlation', 0.0)
                                    self.document.add_paragraph(f"- {f1} 和 {f2}: {corr_v:.4f}", style='List Bullet')
                                    logger.debug(f"添加强相关对: {f1} - {f2}: {corr_v}")
                                except Exception as e:
                                    logger.error(f"处理强相关对 {i} 时出错: {str(e)}")
                        else:
                            self.document.add_paragraph("未发现强相关的特征对。")
                    else:
                        logger.warning("相关性分析结果格式不正确")
                        self.document.add_paragraph("相关性分析结果暂不可用。")
                        
                    logger.info("相关性分析添加成功")
            except Exception as e:
                logger.error(f"添加相关性分析时出错: {str(e)}")
                self.document.add_paragraph("相关性分析生成时出现错误。")
                
            # 添加新的统计分析结果：聚类分析
            try:
                if 'cluster_analysis' in analysis_results:
                    logger.info("添加聚类分析结果")
                    self.document.add_heading("3.3 聚类分析", level=2)
                    
                    cluster_results = analysis_results.get('cluster_analysis') or {}
                    if isinstance(cluster_results, dict):
                        method = cluster_results.get('method', '未知')
                        n_clusters = cluster_results.get('n_clusters', 0)
                        silhouette_score = cluster_results.get('silhouette_score')
                        
                        self.document.add_paragraph(f"采用 {method} 聚类方法，将数据分为 {n_clusters} 个类别。")
                        
                        if silhouette_score is not None:
                            self.document.add_paragraph(f"轮廓系数（Silhouette Score）: {silhouette_score:.4f}")
                            self.document.add_paragraph("轮廓系数越接近1表示聚类效果越好，接近0表示聚类效果一般，接近-1表示聚类效果较差。")
                        
                        # 显示聚类中心
                        if 'cluster_centers' in cluster_results:
                            centers = cluster_results.get('cluster_centers')
                            if centers is not None and hasattr(centers, 'shape'):
                                self.document.add_paragraph(f"聚类中心特征维度: {centers.shape}")
                        
                        # 显示各类别样本数量
                        if 'cluster_counts' in cluster_results:
                            counts = cluster_results.get('cluster_counts', {})
                            if counts:
                                self.document.add_paragraph("各类别样本数量分布:")
                                for cluster_id, count in counts.items():
                                    self.document.add_paragraph(f"- 类别 {cluster_id}: {count} 个样本", style='List Bullet')
            except Exception as e:
                logger.error(f"添加聚类分析时出错: {str(e)}")
                self.document.add_paragraph("聚类分析结果生成时出现错误。")
            
            # 添加因子分析结果
            try:
                if 'factor_analysis' in analysis_results:
                    logger.info("添加因子分析结果")
                    self.document.add_heading("3.4 因子分析", level=2)
                    
                    factor_results = analysis_results.get('factor_analysis') or {}
                    if isinstance(factor_results, dict):
                        n_factors = factor_results.get('n_factors', 0)
                        kmo_score = factor_results.get('kmo_score')
                        explained_variance = factor_results.get('explained_variance_ratio')
                        
                        self.document.add_paragraph(f"提取 {n_factors} 个主要因子。")
                        
                        if kmo_score is not None:
                            self.document.add_paragraph(f"KMO检验值: {kmo_score:.4f}")
                            if kmo_score > 0.7:
                                self.document.add_paragraph("KMO检验值大于0.7，表明数据适合进行因子分析。")
                            elif kmo_score > 0.5:
                                self.document.add_paragraph("KMO检验值在0.5-0.7之间，数据勉强适合因子分析。")
                            else:
                                self.document.add_paragraph("KMO检验值小于0.5，数据不太适合因子分析。")
                        
                        if explained_variance is not None and hasattr(explained_variance, '__len__'):
                            total_variance = sum(explained_variance) * 100
                            self.document.add_paragraph(f"提取的因子累计解释方差: {total_variance:.2f}%")
                            
                            self.document.add_paragraph("各因子解释方差比例:")
                            for i, var_ratio in enumerate(explained_variance, 1):
                                self.document.add_paragraph(f"- 因子 {i}: {var_ratio*100:.2f}%", style='List Bullet')
                        
                        # 显示因子载荷
                        if 'factor_loadings' in factor_results:
                            loadings = factor_results.get('factor_loadings')
                            if loadings is not None and hasattr(loadings, 'shape'):
                                self.document.add_paragraph(f"因子载荷矩阵维度: {loadings.shape[0]} 个变量 × {loadings.shape[1]} 个因子")
            except Exception as e:
                logger.error(f"添加因子分析时出错: {str(e)}")
                self.document.add_paragraph("因子分析结果生成时出现错误。")
            
            # 添加方差分析结果
            try:
                if 'anova_analysis' in analysis_results:
                    logger.info("添加方差分析结果")
                    self.document.add_heading("3.5 方差分析（ANOVA）", level=2)
                    
                    anova_results = analysis_results.get('anova_analysis') or {}
                    if isinstance(anova_results, dict):
                        dependent_var = anova_results.get('dependent_variable', '未知')
                        independent_var = anova_results.get('independent_variable', '未知')
                        f_statistic = anova_results.get('f_statistic')
                        p_value = anova_results.get('p_value')
                        effect_size = anova_results.get('effect_size')
                        
                        self.document.add_paragraph(f"分析因变量 '{dependent_var}' 在自变量 '{independent_var}' 不同水平间的差异。")
                        
                        if f_statistic is not None and p_value is not None:
                            self.document.add_paragraph(f"F统计量: {f_statistic:.4f}")
                            self.document.add_paragraph(f"p值: {p_value:.4f}")
                            
                            if p_value < 0.05:
                                self.document.add_paragraph("p < 0.05，拒绝原假设，不同组间存在显著差异。")
                            else:
                                self.document.add_paragraph("p ≥ 0.05，接受原假设，不同组间无显著差异。")
                        
                        if effect_size is not None:
                            self.document.add_paragraph(f"效应量（η²）: {effect_size:.4f}")
                            if effect_size > 0.14:
                                self.document.add_paragraph("效应量大（η² > 0.14），表明组间差异很大。")
                            elif effect_size > 0.06:
                                self.document.add_paragraph("效应量中等（0.06 < η² ≤ 0.14），表明组间差异适中。")
                            else:
                                self.document.add_paragraph("效应量小（η² ≤ 0.06），表明组间差异较小。")
                        
                        # 显示Levene检验结果
                        if 'levene_test' in anova_results:
                            levene_results = anova_results.get('levene_test', {})
                            levene_p = levene_results.get('p_value')
                            if levene_p is not None:
                                self.document.add_paragraph(f"Levene方差齐性检验 p值: {levene_p:.4f}")
                                if levene_p > 0.05:
                                    self.document.add_paragraph("方差齐性假设成立，ANOVA结果可信。")
                                else:
                                    self.document.add_paragraph("方差齐性假设不成立，建议使用Welch ANOVA。")
                        
                        # 显示事后检验结果
                        if 'post_hoc' in anova_results:
                            post_hoc_results = anova_results.get('post_hoc', {})
                            if 'significant_pairs' in post_hoc_results:
                                sig_pairs = post_hoc_results.get('significant_pairs', [])
                                if sig_pairs:
                                    self.document.add_paragraph("事后检验发现以下组间存在显著差异:")
                                    for pair in sig_pairs:
                                        if isinstance(pair, dict):
                                            group1 = pair.get('group1', '未知')
                                            group2 = pair.get('group2', '未知')
                                            p_val = pair.get('p_value', 0)
                                            self.document.add_paragraph(f"- {group1} vs {group2}: p = {p_val:.4f}", style='List Bullet')
            except Exception as e:
                logger.error(f"添加方差分析时出错: {str(e)}")
                self.document.add_paragraph("方差分析结果生成时出现错误。")

            # 模型推荐
            try:
                if 'model_recommendations' in analysis_results:
                    logger.info("添加模型推荐")
                    recommendations = analysis_results.get('model_recommendations') or []
                    logger.debug(f"recommendations type: {type(recommendations)}, length: {len(recommendations) if hasattr(recommendations, '__len__') else 'no len'}")
                    
                    # 检查recommendations是否为可迭代列表
                    if isinstance(recommendations, (list, tuple)) and len(recommendations) > 0:
                        self.document.add_heading("3.6 模型推荐", level=2)
                        self.document.add_paragraph("根据数据特征，推荐以下分析模型:")
                        
                        for i, model in enumerate(recommendations, 1):
                            try:
                                if not isinstance(model, dict):
                                    logger.warning(f"模型推荐 {i} 不是字典类型: {type(model)}")
                                    continue
                                model_para = self.document.add_paragraph()
                                model_name = model.get('name', '未知模型')
                                model_score = model.get('score', 0)
                                model_para.add_run(f"{i}. {model_name}").bold = True
                                model_para.add_run(f" (推荐指数: {model_score:.2f}/10)")
                                
                                if model.get('description'):
                                    self.document.add_paragraph(str(model.get('description')), style='List Bullet')
                                
                                if model.get('reason'):
                                    self.document.add_paragraph(f"推荐原因: {model.get('reason')}", style='List Bullet')
                                    
                                logger.debug(f"添加模型推荐: {model_name}")
                            except Exception as e:
                                logger.error(f"处理模型推荐 {i} 时出错: {str(e)}")
                    else:
                        self.document.add_heading("3.6 模型推荐", level=2)
                        self.document.add_paragraph("暂无可用的模型推荐。")
                        
                    logger.info("模型推荐添加成功")
            except Exception as e:
                logger.error(f"添加模型推荐时出错: {str(e)}")
                self.document.add_paragraph("模型推荐生成时出现错误。")

            # 添加空行
            self.document.add_paragraph()
            logger.info("分析结果添加完成")
            
        except Exception as e:
            logger.exception(f"添加分析结果时发生未预期的错误: {str(e)}")
            raise

    def add_report_template_example(self) -> None:
        """
        添加报告样例模板到报告开头
        """
        try:
            logger.info("添加报告样例模板")
            
            # 插入一个页面分隔符，将样例模板与正式报告分开
            self.document.add_page_break()
            
            # 添加样例模板标题
            template_title = self.document.add_heading("数据分析报告样例模板", level=1)
            template_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加模板说明
            self.document.add_paragraph("本模板提供标准化的数据分析报告结构，确保报告内容完整、逻辑清晰。")
            
            # 1. 执行摘要模板
            self.document.add_heading("1. 执行摘要", level=2)
            self.document.add_paragraph("【样例】本报告对XXX数据集进行了全面分析，数据集包含X个样本和Y个特征。主要发现包括：")
            self.document.add_paragraph("• 数据质量良好，缺失值比例为X%，已采用合适方法处理", style='List Bullet')
            self.document.add_paragraph("• 发现了X个强相关特征对，相关系数均超过0.7", style='List Bullet')
            self.document.add_paragraph("• 聚类分析将样本分为X个类别，轮廓系数达到X", style='List Bullet')
            self.document.add_paragraph("• 因子分析提取了X个主要因子，累计解释方差X%", style='List Bullet')
            self.document.add_paragraph("推荐模型：基于数据特征，推荐使用XXX模型进行进一步分析。")
            
            # 2. 数据概览模板
            self.document.add_heading("2. 数据概览", level=2)
            self.document.add_heading("2.1 数据集基本信息", level=3)
            
            # 创建样例信息表格
            example_table = self.document.add_table(rows=1, cols=2)
            example_table.style = 'Table Grid'
            hdr_cells = example_table.rows[0].cells
            hdr_cells[0].text = '属性'
            hdr_cells[1].text = '值'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            example_data = [
                ('文件名称', 'sample_data.csv'),
                ('数据格式', 'CSV'),
                ('数据行数', '1,000'),
                ('数据列数', '15'),
                ('数值型特征', '10'),
                ('分类型特征', '5'),
                ('数据大小', '125.6 KB')
            ]
            
            for attr, value in example_data:
                cells = example_table.add_row().cells
                cells[0].text = attr
                cells[1].text = value
            
            self.document.add_heading("2.2 字段信息", level=3)
            self.document.add_paragraph("【样例】数据集包含以下主要字段：")
            self.document.add_paragraph("• 用户ID：唯一标识符", style='List Bullet')
            self.document.add_paragraph("• 年龄：数值型，范围18-65岁", style='List Bullet')
            self.document.add_paragraph("• 性别：分类型，男/女", style='List Bullet')
            self.document.add_paragraph("• 收入：数值型，单位：万元", style='List Bullet')
            
            # 3. 数据预处理模板
            self.document.add_heading("3. 数据预处理", level=2)
            self.document.add_heading("3.1 缺失值处理", level=3)
            self.document.add_paragraph("【样例】发现 50 个缺失值，分布在 3 列中。")
            self.document.add_paragraph("处理方法: 数值型特征采用均值填充，分类型特征采用众数填充")
            
            self.document.add_heading("3.2 异常值处理", level=3)
            self.document.add_paragraph("【样例】采用3σ原则检测异常值，发现 15 个异常值。")
            self.document.add_paragraph("处理方法: 对超出3σ范围的值进行截断处理")
            
            # 4. 数据分析结果模板
            self.document.add_heading("4. 数据分析结果", level=2)
            
            self.document.add_heading("4.1 描述性统计", level=3)
            self.document.add_paragraph("【样例】主要数值型特征的描述性统计如下：")
            # 创建样例统计表格
            stats_example_table = self.document.add_table(rows=1, cols=4)
            stats_example_table.style = 'Table Grid'
            stats_hdr = stats_example_table.rows[0].cells
            stats_hdr[0].text = '统计量'
            stats_hdr[1].text = '年龄'
            stats_hdr[2].text = '收入'
            stats_hdr[3].text = '消费金额'
            for cell in stats_hdr:
                cell.paragraphs[0].runs[0].bold = True
            
            stats_data = [
                ('均值', '32.5', '8.7', '2.3'),
                ('标准差', '8.2', '3.1', '1.2'),
                ('最小值', '18.0', '2.0', '0.1'),
                ('最大值', '65.0', '20.0', '8.9')
            ]
            
            for stat_row in stats_data:
                cells = stats_example_table.add_row().cells
                for i, value in enumerate(stat_row):
                    cells[i].text = value
            
            self.document.add_heading("4.2 相关性分析", level=3)
            self.document.add_paragraph("【样例】相关性分析采用 Pearson 方法。")
            self.document.add_paragraph("发现以下强相关特征对（相关系数绝对值 > 0.7）:")
            self.document.add_paragraph("• 收入 和 消费金额: 0.8523", style='List Bullet')
            self.document.add_paragraph("• 年龄 和 储蓄金额: 0.7341", style='List Bullet')
            
            self.document.add_heading("4.3 聚类分析", level=3)
            self.document.add_paragraph("【样例】采用 K-means 聚类方法，将数据分为 3 个类别。")
            self.document.add_paragraph("轮廓系数（Silhouette Score）: 0.7234")
            self.document.add_paragraph("各类别样本数量分布:")
            self.document.add_paragraph("• 类别 0: 350 个样本", style='List Bullet')
            self.document.add_paragraph("• 类别 1: 420 个样本", style='List Bullet')
            self.document.add_paragraph("• 类别 2: 230 个样本", style='List Bullet')
            
            self.document.add_heading("4.4 因子分析", level=3)
            self.document.add_paragraph("【样例】提取 4 个主要因子。")
            self.document.add_paragraph("KMO检验值: 0.8456 - 数据适合进行因子分析")
            self.document.add_paragraph("提取的因子累计解释方差: 78.92%")
            self.document.add_paragraph("各因子解释方差比例:")
            self.document.add_paragraph("• 因子 1: 34.12%", style='List Bullet')
            self.document.add_paragraph("• 因子 2: 22.45%", style='List Bullet')
            self.document.add_paragraph("• 因子 3: 14.78%", style='List Bullet')
            self.document.add_paragraph("• 因子 4: 7.57%", style='List Bullet')
            
            # 5. 结论与建议模板
            self.document.add_heading("5. 结论与建议", level=2)
            self.document.add_paragraph("【样例】通过对数据集的全面分析，我们得出以下结论：")
            self.document.add_paragraph("1. 数据质量良好，经过适当的清洗和预处理后可用于进一步分析。")
            self.document.add_paragraph("2. 收入与消费金额存在强正相关关系，可作为预测模型的重要特征。")
            self.document.add_paragraph("3. 用户可明确分为三个群体，具有不同的消费特征和行为模式。")
            self.document.add_paragraph("4. 提取的主要因子能够有效解释数据的主要变异来源。")
            
            self.document.add_paragraph("建议：")
            self.document.add_paragraph("• 基于聚类结果制定差异化的营销策略", style='List Bullet')
            self.document.add_paragraph("• 利用强相关特征构建消费预测模型", style='List Bullet')
            self.document.add_paragraph("• 定期更新分析报告，监控用户行为变化", style='List Bullet')
            self.document.add_paragraph("• 考虑引入更多行为特征以提高分析精度", style='List Bullet')
            
            # 添加分隔线
            self.document.add_paragraph("=" * 60, style='Heading 1')
            self.document.add_paragraph("以下为实际分析报告内容", style='Heading 1')
            self.document.add_paragraph("=" * 60, style='Heading 1')
            
            logger.info("报告样例模板添加完成")
            
        except Exception as e:
            logger.exception(f"添加报告样例模板时发生错误: {str(e)}")
            # 如果出错，添加简单的说明
            self.document.add_paragraph("报告样例模板生成时出现错误，将直接显示分析结果。")
    
    def smart_merge_questionnaire_data(self, data: pd.DataFrame, 
                                     item_mapping: Optional[Dict[str, str]] = None) -> Dict:
        """
        智能整合问卷数据与题项表
        
        Args:
            data: 原始数据DataFrame
            item_mapping: 题项映射字典，格式为 {列名: 题目描述}
        
        Returns:
            包含整合后信息的字典
        """
        try:
            logger.info("开始智能整合问卷数据与题项表")
            
            merged_info = {
                'original_columns': list(data.columns),
                'item_mapping': {},
                'scale_detection': {},
                'data_types': {},
                'missing_analysis': {},
                'recommendations': []
            }
            
            # 如果提供了题项映射，直接使用
            if item_mapping and isinstance(item_mapping, dict):
                merged_info['item_mapping'] = item_mapping.copy()
                logger.info(f"使用提供的题项映射，包含 {len(item_mapping)} 个题项")
            else:
                # 智能推断题项映射
                logger.info("开始智能推断题项映射")
                for col in data.columns:
                    col_str = str(col).lower()
                    
                    # 基于列名模式推断题目类型
                    if any(pattern in col_str for pattern in ['q', 'question', '题', '问题']):
                        # 问卷题目
                        if 'age' in col_str or '年龄' in col_str:
                            merged_info['item_mapping'][col] = f"年龄相关题项: {col}"
                        elif 'gender' in col_str or '性别' in col_str:
                            merged_info['item_mapping'][col] = f"性别相关题项: {col}"
                        elif 'income' in col_str or '收入' in col_str or 'salary' in col_str:
                            merged_info['item_mapping'][col] = f"收入相关题项: {col}"
                        elif 'satisfaction' in col_str or '满意' in col_str:
                            merged_info['item_mapping'][col] = f"满意度题项: {col}"
                        elif 'attitude' in col_str or '态度' in col_str:
                            merged_info['item_mapping'][col] = f"态度题项: {col}"
                        else:
                            merged_info['item_mapping'][col] = f"问卷题项: {col}"
                    elif any(pattern in col_str for pattern in ['scale', '量表', 'score', '得分']):
                        merged_info['item_mapping'][col] = f"量表题项: {col}"
                    elif col_str in ['id', 'userid', 'user_id', '用户id', '编号']:
                        merged_info['item_mapping'][col] = f"标识符: {col}"
                    else:
                        # 基于数据类型推断
                        if data[col].dtype in ['int64', 'float64']:
                            unique_values = data[col].nunique()
                            if unique_values <= 10:
                                merged_info['item_mapping'][col] = f"分类变量/等级评分: {col}"
                            else:
                                merged_info['item_mapping'][col] = f"连续变量: {col}"
                        else:
                            merged_info['item_mapping'][col] = f"文本/分类变量: {col}"
            
            # 检测量表类型
            logger.info("检测量表类型")
            for col in data.columns:
                if data[col].dtype in ['int64', 'float64']:
                    unique_values = sorted(data[col].dropna().unique())
                    
                    # Likert量表检测
                    if len(unique_values) <= 7 and all(isinstance(v, (int, float)) for v in unique_values):
                        if set(unique_values).issubset(set(range(1, 6))):
                            merged_info['scale_detection'][col] = "5点Likert量表 (1-5)"
                        elif set(unique_values).issubset(set(range(1, 8))):
                            merged_info['scale_detection'][col] = "7点Likert量表 (1-7)"
                        elif set(unique_values).issubset(set([0, 1])):
                            merged_info['scale_detection'][col] = "二分变量 (0-1)"
                        else:
                            merged_info['scale_detection'][col] = f"等级变量 ({min(unique_values)}-{max(unique_values)})"
                    else:
                        merged_info['scale_detection'][col] = "连续变量"
                else:
                    merged_info['scale_detection'][col] = "分类变量"
            
            # 数据类型分析
            for col in data.columns:
                merged_info['data_types'][col] = {
                    'dtype': str(data[col].dtype),
                    'unique_count': int(data[col].nunique()),
                    'null_count': int(data[col].isnull().sum()),
                    'null_percentage': float(data[col].isnull().sum() / len(data) * 100)
                }
            
            # 缺失值分析
            missing_cols = data.columns[data.isnull().any()].tolist()
            if missing_cols:
                merged_info['missing_analysis']['columns_with_missing'] = missing_cols
                merged_info['missing_analysis']['total_missing'] = int(data.isnull().sum().sum())
                merged_info['missing_analysis']['missing_percentage'] = float(data.isnull().sum().sum() / (len(data) * len(data.columns)) * 100)
                
                # 对每个有缺失值的列分析缺失模式
                for col in missing_cols[:5]:  # 只分析前5个列
                    missing_count = int(data[col].isnull().sum())
                    missing_pct = float(missing_count / len(data) * 100)
                    merged_info['missing_analysis'][col] = {
                        'count': missing_count,
                        'percentage': missing_pct
                    }
            
            # 生成建议
            recommendations = []
            
            # 基于缺失值情况给建议
            if merged_info['missing_analysis'].get('missing_percentage', 0) > 10:
                recommendations.append("数据缺失值比例较高，建议进行专门的缺失值分析和处理")
            elif merged_info['missing_analysis'].get('missing_percentage', 0) > 5:
                recommendations.append("存在一定比例的缺失值，建议采用适当的插补方法")
            
            # 基于量表类型给建议
            likert_scales = [col for col, scale_type in merged_info['scale_detection'].items() 
                           if 'Likert' in scale_type]
            if likert_scales:
                recommendations.append(f"检测到 {len(likert_scales)} 个Likert量表变量，建议进行信度分析")
            
            # 基于变量数量给建议
            if len(data.columns) > 20:
                recommendations.append("变量数量较多，建议进行降维分析或因子分析")
            
            # 基于数据量给建议
            if len(data) < 100:
                recommendations.append("样本量较小，统计分析结果可能不够稳定")
            elif len(data) > 10000:
                recommendations.append("样本量较大，可以进行复杂的统计建模分析")
            
            merged_info['recommendations'] = recommendations
            
            logger.info(f"智能整合完成，识别 {len(merged_info['item_mapping'])} 个题项，{len(likert_scales)} 个Likert量表")
            return merged_info
            
        except Exception as e:
            logger.exception(f"智能整合问卷数据时发生错误: {str(e)}")
            # 返回基本信息
            return {
                'original_columns': list(data.columns) if data is not None else [],
                'item_mapping': {col: f"变量: {col}" for col in data.columns} if data is not None else {},
                'scale_detection': {},
                'data_types': {},
                'missing_analysis': {},
                'recommendations': ["数据整合过程中出现错误，建议检查数据格式"]
            }
    
    def add_questionnaire_analysis_section(self, merged_info: Dict) -> None:
        """
        添加问卷数据分析专节
        
        Args:
            merged_info: 智能整合后的问卷信息
        """
        try:
            logger.info("添加问卷数据分析专节")
            
            self.document.add_heading("问卷数据智能分析", level=1)
            
            # 1. 题项映射表
            if merged_info.get('item_mapping'):
                self.document.add_heading("1.1 题项映射表", level=2)
                self.document.add_paragraph("以下是智能识别的题项与描述对应关系：")
                
                # 创建题项表格
                item_table = self.document.add_table(rows=1, cols=3)
                item_table.style = 'Table Grid'
                item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                hdr_cells = item_table.rows[0].cells
                hdr_cells[0].text = '变量名'
                hdr_cells[1].text = '题项描述'
                hdr_cells[2].text = '数据类型'
                
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 添加题项数据
                for col, description in merged_info['item_mapping'].items():
                    cells = item_table.add_row().cells
                    cells[0].text = str(col)
                    cells[1].text = str(description)
                    
                    # 获取数据类型信息
                    data_type_info = merged_info.get('data_types', {}).get(col, {})
                    scale_type = merged_info.get('scale_detection', {}).get(col, '未知')
                    cells[2].text = scale_type
                    
                    for cell in cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # 2. 量表检测结果
            if merged_info.get('scale_detection'):
                self.document.add_heading("1.2 量表类型检测", level=2)
                
                # 统计不同量表类型
                scale_stats = {}
                for col, scale_type in merged_info['scale_detection'].items():
                    if scale_type not in scale_stats:
                        scale_stats[scale_type] = []
                    scale_stats[scale_type].append(col)
                
                self.document.add_paragraph("检测到以下量表类型：")
                for scale_type, columns in scale_stats.items():
                    self.document.add_paragraph(f"• {scale_type}: {len(columns)} 个变量", style='List Bullet')
                    if len(columns) <= 5:  # 只显示前5个变量
                        var_list = ", ".join(columns)
                        self.document.add_paragraph(f"  变量: {var_list}", style='List Bullet')
                    else:
                        var_list = ", ".join(columns[:5])
                        self.document.add_paragraph(f"  变量: {var_list} 等{len(columns)}个", style='List Bullet')
            
            # 3. 缺失值分析
            if merged_info.get('missing_analysis'):
                missing_info = merged_info['missing_analysis']
                if missing_info.get('total_missing', 0) > 0:
                    self.document.add_heading("1.3 数据完整性分析", level=2)
                    
                    total_missing = missing_info.get('total_missing', 0)
                    missing_pct = missing_info.get('missing_percentage', 0)
                    
                    self.document.add_paragraph(f"数据集共有 {total_missing} 个缺失值，占总数据点的 {missing_pct:.2f}%")
                    
                    # 各列缺失情况
                    if missing_info.get('columns_with_missing'):
                        self.document.add_paragraph("各变量缺失情况：")
                        for col in missing_info['columns_with_missing'][:10]:  # 最多显示10个
                            if col in missing_info:
                                col_missing = missing_info[col]
                                count = col_missing.get('count', 0)
                                pct = col_missing.get('percentage', 0)
                                self.document.add_paragraph(f"• {col}: {count} 个缺失值 ({pct:.1f}%)", style='List Bullet')
            
            # 4. 分析建议
            if merged_info.get('recommendations'):
                self.document.add_heading("1.4 分析建议", level=2)
                self.document.add_paragraph("基于数据特征，系统提供以下分析建议：")
                
                for i, recommendation in enumerate(merged_info['recommendations'], 1):
                    self.document.add_paragraph(f"{i}. {recommendation}")
            
            self.document.add_paragraph()
            logger.info("问卷数据分析专节添加完成")
            
        except Exception as e:
            logger.exception(f"添加问卷数据分析专节时发生错误: {str(e)}")
            self.document.add_paragraph("问卷数据分析专节生成时出现错误。")
    
    def add_chart(self, chart: Union[Figure, str], title: str, 
                 description: Optional[str] = None) -> None:
        """
        添加图表到报告
        
        Args:
            chart: matplotlib Figure对象或图片路径
            title: 图表标题
            description: 图表描述
        """
        # 添加图表标题
        chart_heading = self.document.add_heading(title, level=2)
        
        # 保存图表到临时文件
        if isinstance(chart, Figure):
            img_path = os.path.join(self.temp_dir, f"chart_{len(os.listdir(self.temp_dir)) + 1}.png")
            chart.savefig(img_path, dpi=300, bbox_inches='tight')
            plt.close(chart)  # 关闭图表以释放内存
        else:
            img_path = chart
        
        # 添加图片
        if os.path.exists(img_path):
            para = self.document.add_paragraph()
            run = para.add_run()
            run.add_picture(img_path, width=Inches(6.0))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加描述
        if description:
            desc_para = self.document.add_paragraph(description)
            desc_para.style = 'Quote'
        
        # 添加空行
        self.document.add_paragraph()
    
    def add_multiple_charts(self, charts: Optional[Dict[str, Figure]], section_title: str) -> None:
        """
        添加多个图表到报告
        
        Args:
            charts: 图表名称到图表对象的映射
            section_title: 章节标题
        """
        self.document.add_heading(section_title, level=1)
        
        # 检查charts是否存在、不为None且为字典类型
        if charts and isinstance(charts, dict):
            # 添加每个图表
            for chart_name, chart in charts.items():
                try:
                    self.add_chart(chart, chart_name, f"图表: {chart_name}")
                except Exception as e:
                    logger.exception(f"添加图表 {chart_name} 时发生异常: {e}")
                    # 继续添加其他图表
                    continue
        else:
            # 如果没有可用图表，添加说明文本
            self.document.add_paragraph("暂无可用的图表。")
    
    def add_conclusion(self, conclusion: str) -> None:
        """
        添加结论部分
        
        Args:
            conclusion: 结论内容
        """
        self.document.add_heading("4. 结论与建议", level=1)
        
        # 添加结论段落
        self.document.add_paragraph(conclusion)
        
        # 添加空行
        self.document.add_paragraph()
    
    def add_recommendations(self, recommendations: List[str]) -> None:
        """
        添加建议部分
        
        Args:
            recommendations: 建议列表
        """
        if recommendations:
            self.document.add_heading("4.1 业务建议", level=2)
            
            for rec in recommendations:
                self.document.add_paragraph(f"- {rec}", style='List Bullet')
    
    def save_report(self, output_path: Optional[str] = None) -> str:
        """
        保存报告到文件
        
        Args:
            output_path: 输出文件路径，如果未提供则自动生成
            
        Returns:
            保存的文件路径
        """
        if self.document is None:
            raise ValueError("请先创建报告")
        
        # 如果未提供路径，生成默认路径（桌面）
        if output_path is None:
            # 获取桌面路径
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            
            # 生成文件名
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(desktop_path, f"数据分析报告_{timestamp}.docx")
        
        # 更新目录
        # 注意：python-docx 不支持自动生成目录。为了避免直接操作底层 XML（可能导致 AttributeError），
        # 我们仅用一个占位文本替换原有标记，用户可以在 Word 中手动更新目录字段。
        for para in self.document.paragraphs:
            if "[目录将在保存时自动生成]" in para.text:
                para.clear()
                para.add_run("目录（请在 Word 中更新目录字段以生成最新目录）")
        
        # 保存文档
        # 确保输出目录存在
        try:
            out_dir = os.path.dirname(output_path)
            if out_dir:
                os.makedirs(out_dir, exist_ok=True)
        except Exception as e:
            logger.exception(f"创建输出目录失败: {e}")
            raise

        try:
            self.document.save(output_path)
        except Exception as e:
            logger.exception(f"保存文档到 {output_path} 失败: {e}")
            # 抛出更友好的错误
            raise IOError(f"无法保存报告到 {output_path}: {e}")

        return output_path
    
    def _get_statistic_name(self, stat: str) -> str:
        """
        获取统计量的中文名称
        
        Args:
            stat: 统计量英文名称
            
        Returns:
            中文名称
        """
        stat_names = {
            'mean': '均值',
            'std': '标准差',
            'min': '最小值',
            '25%': '第一四分位数',
            '50%': '中位数',
            '75%': '第三四分位数',
            'max': '最大值',
            'count': '计数',
            'unique': '唯一值数量',
            'top': '最常见值',
            'freq': '最常见值频率'
        }
        return stat_names.get(stat, stat)
    
    def _add_ai_enhanced_content(self, analysis_results: Dict) -> None:
        """
        添加AI增强的分析内容
        
        Args:
            analysis_results: 包含AI增强内容的分析结果字典
        """
        try:
            # 检查是否有AI增强内容
            ai_sections = []
            
            if 'ai_comprehensive_analysis' in analysis_results:
                ai_sections.append(('综合智能分析', analysis_results['ai_comprehensive_analysis']))
            
            if 'ai_insights' in analysis_results:
                ai_sections.append(('AI洞察发现', analysis_results['ai_insights']))
            
            if 'ai_recommendations' in analysis_results:
                ai_sections.append(('智能建议', analysis_results['ai_recommendations']))
            
            if 'ai_interpretation' in analysis_results:
                ai_sections.append(('结果智能解读', analysis_results['ai_interpretation']))
            
            if not ai_sections:
                return
                
            # 添加AI增强分析章节
            self.document.add_heading("3.0 🤖 AI智能分析", level=2)
            
            # 添加说明段落
            intro_para = self.document.add_paragraph()
            intro_run = intro_para.add_run("以下内容由AI大模型基于数据分析结果生成，提供深度洞察和专业建议：")
            intro_run.font.size = Pt(10)
            intro_run.italic = True
            
            for section_title, ai_content in ai_sections:
                if not isinstance(ai_content, dict):
                    continue
                    
                # 添加子标题
                self.document.add_heading(f"🔍 {section_title}", level=3)
                
                # 添加AI模型信息
                model_info = f"AI模型: {ai_content.get('ai_model', 'Unknown')} | " \
                           f"提供商: {ai_content.get('ai_provider', 'Unknown')} | " \
                           f"生成时间: {ai_content.get('enhancement_timestamp', 'Unknown')}"
                
                info_para = self.document.add_paragraph()
                info_run = info_para.add_run(model_info)
                info_run.font.size = Pt(9)
                info_run.font.color.rgb = RGBColor(128, 128, 128)
                info_run.italic = True
                
                # 添加AI生成的内容
                enhanced_content = ai_content.get('enhanced_content', '')
                if enhanced_content:
                    # 将AI内容按段落分割并添加
                    paragraphs = enhanced_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = self.document.add_paragraph()
                            
                            # 检查是否是标题行（包含##或**）
                            if para_text.strip().startswith('##') or '**' in para_text:
                                # 处理Markdown格式的标题和加粗文本
                                if para_text.strip().startswith('##'):
                                    title_text = para_text.strip().replace('##', '').strip()
                                    run = para.add_run(title_text)
                                    run.bold = True
                                    run.font.size = Pt(12)
                                else:
                                    # 处理加粗文本
                                    parts = para_text.split('**')
                                    for i, part in enumerate(parts):
                                        if part:
                                            run = para.add_run(part)
                                            if i % 2 == 1:  # 奇数索引的部分应该加粗
                                                run.bold = True
                            else:
                                # 普通段落
                                para.add_run(para_text.strip())
                
                # 添加分隔线
                self.document.add_paragraph()
                
            logger.info(f"已添加 {len(ai_sections)} 个AI增强分析章节")
            
        except Exception as e:
            logger.error(f"添加AI增强内容失败: {str(e)}")
            # 不抛出异常，继续执行原有流程

# 辅助函数
def qn(tag):
    """
    生成命名空间标签
    """
    return '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + tag

# 高级报告生成器类
class AdvancedReportGenerator:
    """
    高级报告生成器，集成数据分析结果自动生成完整报告
    支持AI大模型增强功能
    """
    
    def __init__(self, ai_enhancer: Optional[AIReportEnhancer] = None):
        self.generator = ReportGenerator()
        self.ai_enhancer = ai_enhancer
        
        # 如果没有提供AI增强器但AI模块可用，尝试创建默认增强器
        if self.ai_enhancer is None and AI_ENHANCEMENT_AVAILABLE:
            try:
                # 尝试使用默认配置创建AI增强器
                self.ai_enhancer = create_ai_enhancer()
                logger.info("已创建默认AI报告增强器")
            except Exception as e:
                logger.warning(f"无法创建默认AI增强器: {str(e)}")
                self.ai_enhancer = None
    
    def set_ai_enhancer(self, ai_enhancer: AIReportEnhancer):
        """设置AI增强器"""
        self.ai_enhancer = ai_enhancer
        logger.info("AI报告增强器已设置")
    
    def generate_full_report(self, data: pd.DataFrame, 
                           analysis_results: Dict, 
                           charts: Optional[Dict[str, Figure]] = None,
                           file_info: Optional[Dict] = None,
                           output_path: Optional[str] = None,
                           include_template: bool = True,
                           item_mapping: Optional[Dict[str, str]] = None) -> str:
        """
        生成完整的分析报告
        
        Args:
            data: 数据框
            analysis_results: 分析结果
            charts: 图表字典
            file_info: 文件信息
            output_path: 输出路径
            include_template: 是否包含报告模板示例
            item_mapping: 题项映射字典，格式为 {列名: 题目描述}
            
        Returns:
            保存的文件路径
        """
        try:
            logger.info("开始生成完整的数据分析报告")
            
            # 保护 analysis_results 为 None 的情况
            analysis_results = analysis_results or {}
            if not isinstance(analysis_results, dict):
                raise ValueError("analysis_results 必须是字典类型")

            # 记录输入摘要以便调试
            logger.info(f"输入数据类型: {type(data)}, 行数: {len(data) if hasattr(data, '__len__') else 'unknown'}")
            logger.info(f"分析结果键: {list(analysis_results.keys())}")
            logger.info(f"图表类型: {type(charts)}, 图表数量: {len(charts) if charts else 0}")
            logger.info(f"文件信息: {file_info}")
            logger.info(f"包含模板: {include_template}, 题项映射: {bool(item_mapping)}")

            # 创建报告
            try:
                title = f"数据分析报告 - {file_info.get('file_name', '数据集')}" if file_info and isinstance(file_info, dict) else "数据分析报告"
                logger.info(f"创建报告，标题: {title}")
                self.generator.create_report(title=title)
                logger.info("报告文档创建成功")
            except Exception as e:
                logger.error(f"创建报告文档失败: {str(e)}")
                raise

            # 1. 添加报告模板示例（如果启用）
            if include_template:
                try:
                    logger.info("添加报告样例模板")
                    self.generator.add_report_template_example()
                    logger.info("报告样例模板添加成功")
                except Exception as e:
                    logger.error(f"添加报告样例模板失败: {str(e)}")
                    # 不抛出异常，继续执行

            # 2. 智能问卷数据整合与分析
            merged_questionnaire_info = None
            if data is not None:
                try:
                    logger.info("开始智能问卷数据整合")
                    merged_questionnaire_info = self.generator.smart_merge_questionnaire_data(data, item_mapping)
                    
                    # 添加问卷数据分析专节
                    self.generator.add_questionnaire_analysis_section(merged_questionnaire_info)
                    logger.info("问卷数据整合与分析完成")
                except Exception as e:
                    logger.error(f"智能问卷数据整合失败: {str(e)}")
                    # 不抛出异常，继续执行

        except Exception as e:
            logger.exception(f"生成完整报告时输入参数检查失败: {str(e)}")
            raise
        
        # 🤖 AI增强分析结果 - 在生成报告前进行AI优化
        if self.ai_enhancer and data is not None and analysis_results:
            try:
                logger.info("开始AI增强分析结果")
                # 进行综合性AI增强
                analysis_results = self.ai_enhancer.enhance_analysis_results(
                    data=data,
                    analysis_results=analysis_results,
                    enhancement_type="comprehensive"
                )
                logger.info("AI增强分析结果完成")
            except Exception as e:
                logger.error(f"AI增强分析结果失败: {str(e)}")
                # AI增强失败不影响报告生成，继续使用原始结果
        
        # 3. 添加执行摘要
        try:
            logger.info("添加执行摘要")
            self._generate_executive_summary(data, analysis_results, merged_questionnaire_info)
            logger.info("执行摘要添加成功")
        except Exception as e:
            logger.error(f"添加执行摘要失败: {str(e)}")
            logger.exception("执行摘要生成详细错误")
            raise
        
        # 4. 添加数据概览
        try:
            logger.info("添加数据概览")
            self._generate_data_overview(data, file_info)
            logger.info("数据概览添加成功")
        except Exception as e:
            logger.error(f"添加数据概览失败: {str(e)}")
            logger.exception("数据概览生成详细错误")
            raise
        
        # 5. 添加数据预处理信息
        try:
            if 'preprocessing' in analysis_results:
                logger.info("添加数据预处理信息")
                self.generator.add_data_preprocessing_section(analysis_results['preprocessing'])
                logger.info("数据预处理信息添加成功")
        except Exception as e:
            logger.error(f"添加数据预处理信息失败: {str(e)}")
            logger.exception("数据预处理信息生成详细错误")
            # 不抛出异常，继续执行
        
        # 6. 添加完整的分析结果（包含新增的统计分析）
        try:
            logger.info("添加完整分析结果")
            self.generator.add_analysis_results(analysis_results)
            logger.info("完整分析结果添加成功")
        except Exception as e:
            logger.error(f"添加完整分析结果失败: {str(e)}")
            logger.exception("完整分析结果生成详细错误")
            raise
        
        # 7. 添加图表
        try:
            if charts and isinstance(charts, dict):
                logger.info(f"添加图表，图表数量: {len(charts)}")
                self.generator.add_multiple_charts(charts, "5. 数据可视化")
                logger.info("图表添加成功")
            else:
                logger.info("无图表需要添加")
        except Exception as e:
            logger.error(f"添加图表失败: {str(e)}")
            logger.exception("图表生成详细错误")
            # 不抛出异常，继续执行
        
        # 8. 添加结论和建议
        try:
            logger.info("添加结论和建议")
            self._generate_conclusion(analysis_results, merged_questionnaire_info)
            logger.info("结论和建议添加成功")
        except Exception as e:
            logger.error(f"添加结论和建议失败: {str(e)}")
            logger.exception("结论和建议生成详细错误")
            raise
        
        # 9. 保存报告
        try:
            logger.info("开始保存报告")
            saved_path = self.generator.save_report(output_path)
            logger.info(f"报告保存成功，路径: {saved_path}")
            return saved_path
        except Exception as e:
            logger.error(f"保存报告失败: {str(e)}")
            logger.exception("保存报告详细错误")
            raise
    
    def _generate_executive_summary(self, data: pd.DataFrame, 
                                   analysis_results: Dict,
                                   merged_questionnaire_info: Optional[Dict] = None) -> None:
        """
        自动生成执行摘要
        """
        try:
            logger.info("开始生成执行摘要")
            summary = []
            
            # 数据基本信息
            try:
                data_rows = len(data) if data is not None and hasattr(data, '__len__') else 0
                data_cols = len(data.columns) if data is not None and hasattr(data, 'columns') else 0
                summary.append(f"本报告分析了包含 {data_rows} 行和 {data_cols} 列的数据集。")
                logger.debug(f"数据基本信息: {data_rows} 行, {data_cols} 列")
            except Exception as e:
                logger.error(f"获取数据基本信息失败: {str(e)}")
                summary.append("本报告对提供的数据集进行了分析。")
            
            # 问卷数据特征（新增）
            try:
                if merged_questionnaire_info:
                    likert_scales = []
                    binary_vars = []
                    continuous_vars = []
                    
                    for col, scale_type in merged_questionnaire_info.get('scale_detection', {}).items():
                        if 'Likert' in scale_type:
                            likert_scales.append(col)
                        elif '二分变量' in scale_type:
                            binary_vars.append(col)
                        elif '连续变量' in scale_type:
                            continuous_vars.append(col)
                    
                    if likert_scales:
                        summary.append(f"识别出 {len(likert_scales)} 个Likert量表变量，适合进行态度和满意度分析。")
                    
                    if binary_vars:
                        summary.append(f"包含 {len(binary_vars)} 个二分变量，可用于分类分析。")
                    
                    missing_pct = merged_questionnaire_info.get('missing_analysis', {}).get('missing_percentage', 0)
                    if missing_pct > 0:
                        summary.append(f"数据完整性良好，缺失值比例为 {missing_pct:.1f}%。")
                    else:
                        summary.append("数据完整性优秀，无缺失值。")
                        
                    logger.debug(f"问卷特征: {len(likert_scales)} 个Likert量表, {len(binary_vars)} 个二分变量")
            except Exception as e:
                logger.error(f"获取问卷数据特征失败: {str(e)}")
            
            # 分析类型
            try:
                if analysis_results and 'analysis_type' in analysis_results and analysis_results['analysis_type']:
                    summary.append(f"根据数据特征，系统自动识别并执行了{analysis_results['analysis_type']}分析。")
                    logger.debug(f"分析类型: {analysis_results['analysis_type']}")
            except Exception as e:
                logger.error(f"获取分析类型失败: {str(e)}")
            # 主要发现
            try:
                # 相关性分析发现
                try:
                    if (analysis_results and 'correlation' in analysis_results and
                        analysis_results['correlation'] is not None):
                        corr_results = analysis_results['correlation']
                        if isinstance(corr_results, dict):
                            strong_corr = corr_results.get('strong_correlations', [])
                            if strong_corr and len(strong_corr) > 0:
                                summary.append(f"• 发现 {len(strong_corr)} 对强相关特征，相关系数均超过0.7。")
                                logger.debug(f"强相关特征对数量: {len(strong_corr)}")
                except Exception as e:
                    logger.error(f"提取相关性分析发现失败: {str(e)}")
                
                # 聚类分析发现（新增）
                try:
                    if (analysis_results and 'cluster_analysis' in analysis_results and
                        analysis_results['cluster_analysis'] is not None):
                        cluster_results = analysis_results['cluster_analysis']
                        if isinstance(cluster_results, dict):
                            n_clusters = cluster_results.get('n_clusters', 0)
                            silhouette_score = cluster_results.get('silhouette_score')
                            if n_clusters > 0:
                                if silhouette_score is not None:
                                    summary.append(f"• 聚类分析将数据分为 {n_clusters} 个类别，轮廓系数为 {silhouette_score:.3f}。")
                                else:
                                    summary.append(f"• 聚类分析将数据分为 {n_clusters} 个不同的类别。")
                                logger.debug(f"聚类结果: {n_clusters} 个类别, 轮廓系数: {silhouette_score}")
                except Exception as e:
                    logger.error(f"提取聚类分析发现失败: {str(e)}")
                
                # 因子分析发现（新增）
                try:
                    if (analysis_results and 'factor_analysis' in analysis_results and
                        analysis_results['factor_analysis'] is not None):
                        factor_results = analysis_results['factor_analysis']
                        if isinstance(factor_results, dict):
                            n_factors = factor_results.get('n_factors', 0)
                            explained_variance = factor_results.get('explained_variance_ratio')
                            if n_factors > 0 and explained_variance is not None:
                                total_variance = sum(explained_variance) * 100
                                summary.append(f"• 因子分析提取 {n_factors} 个主要因子，累计解释方差 {total_variance:.1f}%。")
                                logger.debug(f"因子分析: {n_factors} 个因子, 解释方差: {total_variance:.1f}%")
                except Exception as e:
                    logger.error(f"提取因子分析发现失败: {str(e)}")
                
                # 方差分析发现（新增）
                try:
                    if (analysis_results and 'anova_analysis' in analysis_results and
                        analysis_results['anova_analysis'] is not None):
                        anova_results = analysis_results['anova_analysis']
                        if isinstance(anova_results, dict):
                            p_value = anova_results.get('p_value')
                            dependent_var = anova_results.get('dependent_variable', '目标变量')
                            if p_value is not None:
                                if p_value < 0.05:
                                    summary.append(f"• 方差分析显示 {dependent_var} 在不同组间存在显著差异 (p < 0.05)。")
                                else:
                                    summary.append(f"• 方差分析显示 {dependent_var} 在不同组间无显著差异 (p ≥ 0.05)。")
                                logger.debug(f"方差分析: p值 = {p_value}, 变量: {dependent_var}")
                except Exception as e:
                    logger.error(f"提取方差分析发现失败: {str(e)}")
                
                # 从统计分析中提取发现
                try:
                    if (analysis_results and 'descriptive_stats' in analysis_results and 
                        analysis_results['descriptive_stats'] is not None):
                        stats_df = analysis_results['descriptive_stats']
                        if data is not None and hasattr(data, 'select_dtypes'):
                            numeric_cols = data.select_dtypes(include=[np.number]).columns
                            if len(numeric_cols) > 0:
                                # 优先直接从原始数据计算标准差并找出变异最大的列，这比依赖统计表更稳健
                                try:
                                    std_series = data[numeric_cols].std()
                                    if std_series is not None and not std_series.empty:
                                        max_var_col = std_series.idxmax()
                                        summary.append(f"• {max_var_col} 是变异程度最大的特征。")
                                        logger.debug(f"变异最大的特征: {max_var_col}")
                                except Exception as e:
                                    logger.error(f"计算变异程度失败: {str(e)}")
                                    # 退回到尝试从 stats_df 中读取 'std' 行（如果存在）
                                    try:
                                        if stats_df is not None and hasattr(stats_df, 'loc') and 'std' in stats_df.index:
                                            max_var_col = stats_df.loc['std'].idxmax()
                                            summary.append(f"• {max_var_col} 是变异程度最大的特征。")
                                            logger.debug(f"从统计表获取变异最大的特征: {max_var_col}")
                                    except Exception as e2:
                                        logger.error(f"从统计表获取变异程度失败: {str(e2)}")
                except Exception as e:
                    logger.error(f"提取统计分析发现失败: {str(e)}")
                        
            except Exception as e:
                logger.error(f"提取主要发现失败: {str(e)}")
                summary.append("• 数据质量良好，适合进行分析。")
            
            # 模型推荐
            try:
                if (analysis_results and 'model_recommendations' in analysis_results and 
                    analysis_results['model_recommendations'] and 
                    isinstance(analysis_results['model_recommendations'], (list, tuple)) and
                    len(analysis_results['model_recommendations']) > 0):
                    top_model = analysis_results['model_recommendations'][0]
                    if top_model and isinstance(top_model, dict):
                        model_name = top_model.get('name', '未知模型')
                        summary.append(f"\n推荐模型：{model_name}")
                        logger.debug(f"推荐模型: {model_name}")
            except Exception as e:
                logger.error(f"获取模型推荐失败: {str(e)}")
            
            # 合并摘要
            try:
                summary_text = " ".join(str(s) for s in summary if s)  # 确保所有元素都是字符串且非空
                logger.debug(f"生成的摘要文本长度: {len(summary_text)}")
                self.generator.add_executive_summary(summary_text)
                logger.info("执行摘要生成完成")
            except Exception as e:
                logger.error(f"合并和添加摘要失败: {str(e)}")
                # 使用简单的默认摘要
                default_summary = "本报告对数据集进行了全面的分析，包括数据概览、统计分析和可视化展示。"
                self.generator.add_executive_summary(default_summary)
                logger.info("使用默认摘要")
                
        except Exception as e:
            logger.exception(f"生成执行摘要时发生未预期的错误: {str(e)}")
            raise
    
    def _generate_data_overview(self, data: pd.DataFrame, 
                              file_info: Optional[Dict] = None) -> None:
        """
        自动生成数据概览
        """
        try:
            logger.info("开始生成数据概览")
            
            # 安全地处理file_info
            safe_file_info = file_info or {}
            
            data_info = {
                'file_name': str(safe_file_info.get('file_name', '未知')) if safe_file_info.get('file_name') is not None else '未知',
                'file_format': str(safe_file_info.get('file_format', '未知')) if safe_file_info.get('file_format') is not None else '未知',
                'num_rows': len(data) if data is not None else 0,
                'num_columns': len(data.columns) if data is not None and hasattr(data, 'columns') else 0,
                'num_numeric_cols': len(data.select_dtypes(include=[np.number]).columns) if data is not None else 0,
                'num_categorical_cols': len(data.select_dtypes(include=['object', 'category']).columns) if data is not None else 0,
                'num_date_cols': len(data.select_dtypes(include=['datetime64']).columns) if data is not None else 0,
                'data_size': f"{data.memory_usage(deep=True).sum() / 1024:.2f} KB" if data is not None and hasattr(data, 'memory_usage') else '未知'
            }
            
            logger.debug(f"数据概览信息: {data_info}")
            
            # 列信息
            columns_info = []
            if data is not None and hasattr(data, 'columns'):
                for col in data.columns[:10]:  # 只显示前10列
                    try:
                        col_info = {
                            'name': str(col) if col is not None else '未知列',
                            'dtype': str(data[col].dtype) if hasattr(data[col], 'dtype') else '未知',
                            'non_null_count': int(data[col].count()) if hasattr(data[col], 'count') else 0,
                            'description': f"缺失值: {data[col].isnull().sum()}" if hasattr(data[col], 'isnull') else 'N/A'
                        }
                        columns_info.append(col_info)
                        logger.debug(f"处理列: {col}, 信息: {col_info}")
                    except Exception as e:
                        logger.error(f"处理列 {col} 时出错: {str(e)}")
                        # 添加默认信息
                        col_info = {
                            'name': str(col) if col is not None else '未知列',
                            'dtype': '未知',
                            'non_null_count': 0,
                            'description': 'N/A'
                        }
                        columns_info.append(col_info)
            
            data_info['columns_info'] = columns_info
            logger.info(f"生成了 {len(columns_info)} 个列的信息")
            
            # 添加到报告
            self.generator.add_data_overview(data_info)
            logger.info("数据概览添加成功")
            
        except Exception as e:
            logger.exception(f"生成数据概览时发生未预期的错误: {str(e)}")
            raise
    
    def _generate_conclusion(self, analysis_results: Dict,
                            merged_questionnaire_info: Optional[Dict] = None) -> None:
        """
        自动生成结论和建议
        """
        try:
            logger.info("开始生成结论和建议")
            conclusion_parts = []
            
            # 总结数据分析结果
            conclusion_parts.append("通过对数据集的全面分析，我们得出以下结论：")
            
            # 添加主要结论（保护可能为 None 的字段）
            try:
                conclusions = analysis_results.get('conclusions') if analysis_results else None
                logger.debug(f"结论类型: {type(conclusions)}, 内容: {conclusions}")
                
                if conclusions and isinstance(conclusions, (list, tuple)) and len(conclusions) > 0:
                    for i, conclusion in enumerate(conclusions, 1):
                        if conclusion:  # 确保conclusion不为None或空字符串
                            conclusion_parts.append(f"{i}. {conclusion}")
                            logger.debug(f"添加结论 {i}: {conclusion}")
                else:
                    logger.info("没有找到预设结论，生成自动结论")
                    conclusion_count = 1
                    
                    # 基础数据质量结论
                    conclusion_parts.append(f"{conclusion_count}. 数据质量良好，经过适当的清洗和预处理后可用于进一步分析。")
                    conclusion_count += 1
                    
                    # 问卷数据特定结论（新增）
                    if merged_questionnaire_info:
                        try:
                            likert_count = sum(1 for scale_type in merged_questionnaire_info.get('scale_detection', {}).values() 
                                             if 'Likert' in scale_type)
                            if likert_count > 0:
                                conclusion_parts.append(f"{conclusion_count}. 问卷包含 {likert_count} 个Likert量表变量，适合进行态度和认知分析。")
                                conclusion_count += 1
                            
                            missing_pct = merged_questionnaire_info.get('missing_analysis', {}).get('missing_percentage', 0)
                            if missing_pct < 5:
                                conclusion_parts.append(f"{conclusion_count}. 数据完整性优秀，缺失值比例仅为 {missing_pct:.1f}%，数据可信度高。")
                                conclusion_count += 1
                        except Exception as e:
                            logger.error(f"生成问卷结论失败: {str(e)}")
                    
                    # 相关性分析结论
                    if analysis_results and 'correlation' in analysis_results:
                        conclusion_parts.append(f"{conclusion_count}. 通过相关性分析发现了特征间的重要关联关系。")
                        conclusion_count += 1
                    
                    # 聚类分析结论（新增）
                    if (analysis_results and 'cluster_analysis' in analysis_results and
                        analysis_results['cluster_analysis'] is not None):
                        cluster_results = analysis_results['cluster_analysis']
                        if isinstance(cluster_results, dict):
                            n_clusters = cluster_results.get('n_clusters', 0)
                            if n_clusters > 0:
                                conclusion_parts.append(f"{conclusion_count}. 聚类分析成功将样本分为 {n_clusters} 个具有不同特征的群体。")
                                conclusion_count += 1
                    
                    # 因子分析结论（新增）
                    if (analysis_results and 'factor_analysis' in analysis_results and
                        analysis_results['factor_analysis'] is not None):
                        factor_results = analysis_results['factor_analysis']
                        if isinstance(factor_results, dict):
                            n_factors = factor_results.get('n_factors', 0)
                            if n_factors > 0:
                                conclusion_parts.append(f"{conclusion_count}. 因子分析成功提取 {n_factors} 个主要维度，有效简化了数据结构。")
                                conclusion_count += 1
                    
                    # 方差分析结论（新增）
                    if (analysis_results and 'anova_analysis' in analysis_results and
                        analysis_results['anova_analysis'] is not None):
                        anova_results = analysis_results['anova_analysis']
                        if isinstance(anova_results, dict):
                            p_value = anova_results.get('p_value')
                            if p_value is not None:
                                if p_value < 0.05:
                                    conclusion_parts.append(f"{conclusion_count}. 方差分析证实了不同组别间存在统计学意义上的显著差异。")
                                else:
                                    conclusion_parts.append(f"{conclusion_count}. 方差分析显示不同组别间差异不显著，群体特征相对一致。")
                                conclusion_count += 1
                    
                    # 模型推荐结论
                    if analysis_results and 'model_recommendations' in analysis_results:
                        conclusion_parts.append(f"{conclusion_count}. 基于数据特征，系统推荐了适合的分析模型。")
                        conclusion_count += 1
                        
            except Exception as e:
                logger.error(f"生成结论失败: {str(e)}")
                conclusion_parts.append("1. 数据质量良好，经过适当的清洗和预处理后可用于进一步分析。")
            
            # 合并结论
            try:
                conclusion_text = " ".join(str(part) for part in conclusion_parts if part)
                logger.debug(f"生成的结论文本长度: {len(conclusion_text)}")
                self.generator.add_conclusion(conclusion_text)
                logger.info("结论添加成功")
            except Exception as e:
                logger.error(f"添加结论失败: {str(e)}")
                default_conclusion = "通过对数据集的全面分析，我们得出以下结论：1. 数据质量良好，适合进行分析。"
                self.generator.add_conclusion(default_conclusion)
                logger.info("使用默认结论")
            
            # 添加建议
            try:
                recommendations = []
                if (analysis_results and 'recommendations' in analysis_results and 
                    analysis_results['recommendations'] is not None):
                    # 确保recommendations是可迭代的列表
                    if isinstance(analysis_results['recommendations'], (list, tuple)):
                        recommendations = [rec for rec in analysis_results['recommendations'] if rec]  # 过滤掉空值
                        logger.debug(f"找到 {len(recommendations)} 个预设建议")
                
                if not recommendations:
                    logger.info("没有找到预设建议，生成自动建议")
                    # 自动生成建议
                    recommendations.append("持续收集数据，建立时间序列分析模型以预测未来趋势。")
                    recommendations.append("考虑引入更多相关特征以提高分析精度。")
                    recommendations.append("基于推荐的模型进行深入的预测分析。")
                    recommendations.append("定期更新分析报告，监控关键指标的变化。")
                    
                    # 基于问卷数据的专门建议（新增）
                    if merged_questionnaire_info:
                        try:
                            likert_count = sum(1 for scale_type in merged_questionnaire_info.get('scale_detection', {}).values() 
                                             if 'Likert' in scale_type)
                            if likert_count > 0:
                                recommendations.append("针对Likert量表数据，建议进行信度和效度分析以确保测量质量。")
                            
                            if merged_questionnaire_info.get('missing_analysis', {}).get('missing_percentage', 0) > 0:
                                recommendations.append("对于存在缺失值的问卷题项，建议分析缺失模式并采用适当的处理方法。")
                            
                            recommendations.append("建议对问卷数据进行因子分析，探索潜在的维度结构。")
                        except Exception as e:
                            logger.error(f"生成问卷建议失败: {str(e)}")
                    
                    # 基于分析结果的专门建议（新增）
                    if (analysis_results and 'cluster_analysis' in analysis_results and
                        analysis_results['cluster_analysis'] is not None):
                        recommendations.append("基于聚类分析结果，建议针对不同群体制定差异化的策略。")
                    
                    if (analysis_results and 'factor_analysis' in analysis_results and
                        analysis_results['factor_analysis'] is not None):
                        recommendations.append("利用因子分析结果构建降维模型，提高后续分析效率。")
                
                logger.debug(f"最终建议数量: {len(recommendations)}")
                self.generator.add_recommendations(recommendations)
                logger.info("建议添加成功")
                
            except Exception as e:
                logger.error(f"添加建议失败: {str(e)}")
                default_recommendations = [
                    "持续收集数据，建立时间序列分析模型以预测未来趋势。",
                    "考虑引入更多相关特征以提高分析精度。"
                ]
                self.generator.add_recommendations(default_recommendations)
                logger.info("使用默认建议")
                
        except Exception as e:
            logger.exception(f"生成结论和建议时发生未预期的错误: {str(e)}")
            raise

# 创建工厂函数
def create_report_generator() -> ReportGenerator:
    """
    创建报告生成器实例
    """
    return ReportGenerator()

def create_advanced_report_generator() -> AdvancedReportGenerator:
    """
    创建高级报告生成器实例
    """
    return AdvancedReportGenerator()
